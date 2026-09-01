# -*- coding: utf-8 -*-
"""Build E-Acceptance requirement doc v20260901 based on user's 20260831 version."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ---- base styles ----
st = doc.styles['Normal']
st.font.name = 'DengXian'
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'DengXian')

for i, sz in [(1, 20), (2, 15), (3, 12.5), (4, 11)]:
    s = doc.styles['Heading %d' % i]
    s.font.name = 'DengXian'
    s.font.size = Pt(sz)
    s.font.color.rgb = RGBColor(0x1F, 0x3B, 0x63)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), 'DengXian')


def H(lvl, text):
    doc.add_heading(text, level=lvl)


def P(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def BULLET(text):
    doc.add_paragraph(text, style='List Bullet')


def NUM(text):
    doc.add_paragraph(text, style='List Number')


def TODO(code, text):
    """Open item marker - rendered in orange bold so it's findable in Word."""
    p = doc.add_paragraph()
    r = p.add_run('【待确认 %s】' % code)
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
    return p


def NEW(text):
    """Mark content newly added vs 20260831 version."""
    p = doc.add_paragraph()
    r = p.add_run('【本次新增】')
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x70, 0x30)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0x00, 0x70, 0x30)
    return p


def FIX(text):
    p = doc.add_paragraph()
    r = p.add_run('【本次修正】')
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    return p


def T(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(str(h))
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


# =====================================================================
# 封面
# =====================================================================
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('E-Acceptance（电子签约）需求文档')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x63)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('版本 v3.0　　2026-09-01')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()
T(['项目', '内容'], [
    ['文档版本', 'v3.0（基线：v2.0；原始基线 E-Acceptance需求20260831.docx）'],
    ['对照基准', 'HP075 FRS v5.3、HP047 FRS v1.7、Auto Loan eAcceptance Training Deck 2022.09'],
    ['本次变更来源', '2026-09-01 B3 签约需求评审（任晓宇、张骏、姚佳宇、游路、游义明）'],
    ['本次主要变更', '新增签约模式切换规则章节；CILT 操作三步合并为两步；OTP 归属更正为 RIB 侧；归档位置分工明确；首次 Manual 的 CILT 免重签'],
    ['标记说明', '绿色【本次新增】= 新补充内容；红色【本次修正】= 修正既有错误；橙色【待确认 P-xx】= 需业务或客户确认'],
], widths=[3.2, 12.3])

doc.add_page_break()

# =====================================================================
# 0 修订说明
# =====================================================================
H(1, '0. 本版修订说明')

H(2, '0.1 补充的整块功能（原始FRS有、v1.0未覆盖）')
T(['编号', '功能', '原文出处', '本版位置'], [
    ['A-1', '放款补件自动上传（Defect Rectification）', 'HP075 第6章 / Training Deck 2.3', '第10.5节'],
    ['A-2', 'File Attachment 签约类型双向打标（四场景）', 'HP075 5.11 / Training Deck 2.4', '第10.4节'],
    ['A-3', '定时任务改造，含 LHDN 印花税申报文件拆分', 'HP075 第11章', '第15章'],
    ['A-4', '五个列表/查询页字段增强 + CRA两个界面四字段', 'HP075 5.1–5.5', '第9.1–9.2节'],
    ['A-5', 'E-Acceptance 按钮权限与显示规则', 'HP075 5.8', '第9.3节'],
], widths=[1.5, 6.0, 4.5, 3.5])

H(2, '0.2 补充的关键规则')
T(['编号', '规则', '原文出处', '本版位置'], [
    ['B-1', '四个流程级 General Constant 开关', 'HP075 13.1', '第14.1节'],
    ['B-2', '工作流校验器清单与用户可见错误文案', 'HP075 12.3 / HP047 第6章', '第13章'],
    ['B-3', 'Amendment 变更追踪 46 个字段完整清单', 'HP075 7.1', '第7.4节'],
    ['B-4', '8个 Letter Template + 签名替代规则 + 2条排版合规规则', 'HP075 13.2.1 / 13.4', '第14章'],
    ['B-5', 'hpAgreementInd 复合状态值 + 四个盖章指示器', 'HP075 10.2 / 12.1.1', '第16.3节'],
    ['B-6', '生物识别准入/强制矩阵可配置性 + Owner 角色', 'HP047 第5章', '第8.2节'],
], widths=[1.5, 6.0, 4.5, 3.5])

H(2, '0.3 修正的内部矛盾')
T(['编号', '问题', '处理'], [
    ['C-1', 'Hybrid 流程 5 张表共约 12 处 Mode 误写为 E-Acceptance', '统一修正为 Mode=Manual Acceptance'],
    ['C-2', '术语表称 E-Acceptance(Hybrid)，正文称 Manual Acceptance(Hybrid)', '统一为 Manual Acceptance (Hybrid)，并在术语表注明与原FRS的归类差异'],
    ['C-3', '指纹重试次数在 4 处分别写作 3 次与 999 次', '统一为「默认3次，由 BIOMETRIC_ATTEMPT_COUNTER 配置，系统上限999」'],
    ['C-4', '章节编号重复（两个第12章）、11.2.5 层级错位', '全文重新编号'],
    ['C-5', 'Fail 收进 Pending 后，列表页无法区分「未开始」与「已失败」', '列表页增加失败态展示文本，见第3.3节'],
    ['C-6', 'Print Log 字段重设计，丢失原文四列 Y/Blank 文档接受标记', '恢复四列标记并与新增字段并存，见第11.3节'],
    ['C-7', '客户反馈第2/5/11条未落入文档', '分别补入第8.1节、第4.3节Step2、第11.3节'],
], widths=[1.5, 7.0, 7.0])

H(2, '0.4 v3.0 依据 2026-09-01 评审的变更')
T(['编号', '变更', '本版位置'], [
    ['R-1', '新增「签约模式切换规则」独立章节，统一汇总原散落各处的切换条件', '第4章'],
    ['R-2', 'CILT 操作由三个按钮（同步 / 发短信 / 刷新）合并为两步；明确失败处理与断点', '8.3 节'],
    ['R-3', 'e-Tracker 同步失败时停留 Step 1、不发短信、不记日志；仅记录短信发送状态', '8.3 节'],
    ['R-4', 'Refresh 刷新范围为整屏（含短信状态）；进入页面可自动刷新；Resend 后自动刷新', '8.3 节'],
    ['R-5', '客户在 e-Tracker 确认后系统不自动推进，需 Sales 主动 Refresh', '8.3 节'],
    ['R-6', '新增「首次签约为 Manual 时 CILT 免重签」规则', '8.7 节'],
    ['R-7', 'OTP 全流程归属更正为 RIB 侧；CrediOS 仅通知与等待，不实现 OTP 界面', '5.3 Step 6、9.5 节'],
    ['R-8', '明确 Document Library（全版本 + Template/Signed 标识）与文件中心（仅最新）的分工', '11.1 节'],
    ['R-9', '不符合 E-Acceptance 资格时**不显示**该入口（非置灰）', '2.3 节、4.1 节'],
    ['R-10', '强制转 Manual 后 E-Acceptance 本轮锁定，须后端持久化', '4.2 节、4.5 节'],
    ['R-11', '指纹与 OTP 失败转 Manual 的次数补充「可配置」说明', '5.3 Step 3、14.2 节'],
    ['R-12', '新增「浏览器刷新」场景：刷新后停留当前业务步骤', '5.5 节'],
    ['R-13', 'Back 按钮规则细化：流程步骤中保留，CILT/完成态移除', '5.5 节'],
], widths=[1.5, 9.5, 4.5])

H(2, '0.5 已结案的原待确认项')
T(['原编号', '事项', '结论'], [
    ['P-07', 'Amendment 变更追踪的完整字段清单', '原始FRS HP075 p.33-34 已明确定义 46 个字段，见第7.4节，无需再确认'],
    ['P-22', 'Letter Template 完整清单', '原始FRS HP075 13.2.1 已定义 8 个模板代码及字段替代规则，见第14章，无需再确认'],
    ['P-11', 'CILT Appeal 流程中的 E-Acceptance 处理逻辑', '原始FRS 4.1.4 已明确：签约前 CILT 退回按新案件处理，签约后按 Amendment 处理，见第7.1节'],
], widths=[1.8, 6.0, 7.7])

doc.add_page_break()

# =====================================================================
# 1 术语体系
# =====================================================================
H(1, '1. 术语体系')

H(2, '1.1 核心层级关系')
P('Acceptance 是流程层级概念，Signing 是动作层级概念。二者不可混用：')
BULLET('E-Acceptance = 电子签约流程（整体）')
BULLET('E-signing = 电子签字（动作）')
BULLET('Paper Signing = 纸质签字（动作）')

T(['层级', '术语', '含义', '说明'], [
    ['流程层级', 'Acceptance', '签约/接受', '指整个签约流程，涵盖从身份验证到签署完成的端到端过程'],
    ['动作层级', 'Signing', '签字/签署', '指 Acceptance 流程中的签字动作环节，是 Acceptance 的子步骤'],
], widths=[2.2, 3.0, 2.5, 7.8])

H(2, '1.2 统一术语表')
T(['术语', '英文', '层级', '定义'], [
    ['电子签约', 'E-Acceptance', '流程', 'Biometric Verification + E-signing，全程线上完成'],
    ['人工签约（混合）', 'Manual Acceptance (Hybrid)', '流程', 'Biometric Verification + Paper Signing，认证在线、签署线下'],
    ['人工签约', 'Manual Acceptance', '流程', '纸质签字流程，无在线身份验证环节'],
    ['签字/签署', 'Signing', '动作', '客户在协议文件上签字确认的动作'],
    ['电子签字', 'E-signing', '动作', '客户在 RIB 网银门户上以电子方式签署协议文件'],
    ['纸质签字', 'Paper Signing', '动作', '客户在纸质协议文件上亲笔签字'],
    ['生物识别验证', 'Biometric Verification', '动作', '通过指纹扫描验证客户身份的环节'],
    ['身份验证', 'Identity Verification', '动作', '第1次 Biometric，用于验证客户身份'],
    ['签约确认', 'E-Acceptance Confirmation', '动作', '第2次 Biometric，用于确认客户签约意愿'],
    ['一次性密码', 'OTP (One-Time Password)', '数据', '6位数字验证码，作为第2次 Biometric 的备选验证方式'],
    ['主申请人', 'Primary Applicant', '角色', '贷款申请人本人'],
    ['担保人', 'Guarantor', '角色', '为贷款提供担保的第三方'],
    ['车主/物权人', 'Owner', '角色', '车辆登记所有人，可能不同于借款人；企业客群下需单独做生物识别'],
    ['卖家（直客）', 'Seller (Direct)', '角色', 'Direct 渠道下的卖方，是否需生物识别由后台开关控制'],
    ['经销商', 'Dealer', '角色', '汽车经销商，贷款渠道'],
], widths=[2.6, 4.0, 1.5, 7.4])

FIX('术语归类调整说明：原始 FRS 将「Biometric + Paper Signing」称为 E-Acceptance (Hybrid)，归入 E-Acceptance 层级。本文档将其重命名为 Manual Acceptance (Hybrid) 并归入 Manual 层级，理由是该路径不发生任何电子签署动作。客户对照原文时需注意此差异。')

H(2, '1.3 系统术语表')
T(['术语', '全称', '说明'], [
    ['CrediOS', '—', '本次改造的目标系统，承载签约全流程'],
    ['LOAD$', 'Loan Administration System', 'HLB 旧版贷款处理系统，与 CrediOS 交互'],
    ['RIB', 'Retail Internet Banking', '零售网银门户，客户在此完成 E-signing'],
    ['eTracker', 'Electronic Tracker', '客户签约状态查询平台，Amendment 远程重签也在此完成'],
    ['EAI', 'Enterprise Application Integration', '企业应用集成中间件'],
    ['AEM', 'Adobe Experience Manager', 'eTracker 的底层平台'],
    ['RBS', 'Retail Banking System', '零售银行系统'],
    ['EDMS', '—', '文件归档管理系统'],
    ['Infobip', '—', '短信网关，用于 Amendment 通知与 OTP 下发'],
    ['LHDN', 'Lembaga Hasil Dalam Negeri', '马来西亚税务局，印花税申报接收方'],
], widths=[2.6, 5.0, 7.9])

H(2, '1.4 关联系统交互关系')
T(['交互关系', '说明'], [
    ['CrediOS ↔ Biometric Device', 'CrediOS 客户端通过驱动调用生物识别设备，设备返回识别结果'],
    ['CrediOS ↔ EAI', 'CrediOS 通过 EAI 中间件与 RIB、eTracker 通信'],
    ['EAI ↔ RIB', 'EAI 转发签约邀请至 RIB；RIB 回传 Acceptance Flag'],
    ['EAI ↔ eTracker', 'CrediOS 通过 EAI 向 eTracker 提供状态数据；eTracker 响应客户查询'],
    ['CrediOS → Infobip', 'CrediOS 通过 Infobip 短信网关发送 OTP 和签约通知'],
    ['CrediOS ↔ HOST', 'CrediOS 与旧版 LOAD$ 核心系统进行数据同步'],
    ['CrediOS ↔ EDMS', 'CrediOS 调用 EDMS 进行文件归档与查询'],
    ['CrediOS → LHDN', '按签约类型分别生成印花税申报 XML（见第15章）'],
], widths=[5.0, 10.5])

H(2, '1.5 生物识别硬件操作流程')
P('身份验证由销售人员在系统中操作，包含以下步骤：')
NUM('启动扫描：销售人员在系统中找到该订单的"身份验证"项，点击「Start Biometric Scan」')
NUM('读取身份证：当系统提示读卡器已就绪（Reader is ready）时，将客户的身份证（MyKad）插入设备，点击「Read MyKad」')
NUM('验证指纹：系统界面提示客户"灯亮时请将拇指放在扫描仪上"，客户此时按下指纹')
NUM('完成：指纹与身份证信息匹配成功后，系统显示"Verified OK. Please Remove Card."，此时移开手指并拔出身份证')

doc.add_page_break()

# =====================================================================
# 2 签约类型与客群分流
# =====================================================================
H(1, '2. 签约类型与客群分流')

H(2, '2.1 签约类型总览')
T(['签约类型', '身份验证', '签署方式', '适用客群'], [
    ['E-Acceptance (Full)', 'Biometric Verification', 'E-signing（RIB网银在线签署）', 'I-1：Dealer + 无担保人 + 本地'],
    ['Manual Acceptance (Hybrid)', 'Biometric Verification', 'Paper Signing（纸质签字）', 'I-3 / I-4 / I-6：有担保人 或 Direct渠道'],
    ['Manual Acceptance', '无在线验证（纸质身份核对）', 'Paper Signing（纸质签字）', 'I-2 / I-5 / Non-Individual：外籍 或 非个人'],
], widths=[4.0, 3.6, 4.0, 3.9])

P('层级说明：Manual Acceptance (Hybrid) 与 Manual Acceptance 同属 Manual Acceptance 流程层级，前端展示时均显示为「Manual Acceptance」，名称上不做区隔，仅在流程推进时分化。客户在前端看到没有差异。')

H(2, '2.2 分客群签约类型明细表')
T(['客群', '渠道', '担保', '国籍', 'Biometric', 'Signing方式', 'Acceptance类型'], [
    ['Individual', 'Dealer', '无', '本地（ID1 Red/Blue）', '✅ 主申请人', 'E-signing', 'E-Acceptance'],
    ['Individual', 'Dealer', '无', '外籍（Passport）', '❌ 豁免', 'Paper Signing', 'Manual Acceptance'],
    ['Individual', 'Dealer', '有', '任意', '✅ 主申请人+担保人', 'Paper Signing', 'Manual Acceptance (Hybrid)'],
    ['Individual', 'Direct', '无', '本地（ID1 Red/Blue）', '✅ 主申请人', 'Paper Signing', 'Manual Acceptance (Hybrid)'],
    ['Individual', 'Direct', '无', '外籍', '❌ 豁免', 'Paper Signing', 'Manual Acceptance'],
    ['Individual', 'Direct', '有', '任意', '✅ 主申请人+担保人', 'Paper Signing', 'Manual Acceptance (Hybrid)'],
    ['Non-Individual', '—', '—', '—', '见 2.4 节', 'Paper Signing', 'Manual Acceptance'],
], widths=[2.3, 1.5, 1.2, 3.0, 2.6, 2.2, 2.7])

H(2, '2.3 E-Acceptance 准入条件')
T(['条件', '说明'], [
    ['客户类型', 'Individual（个人申请）'],
    ['渠道', 'Dealer（汽车经销商渠道，含 Panel Dealer 和 Non-panel Dealer）'],
    ['担保人', '无担保人'],
    ['国籍/证件', '本地（ID1 Red / ID1 Blue）'],
], widths=[3.0, 12.5])

P('系统通过 E-Acceptance Flag 字段判定，四项条件须同时满足。判定结果同时受第14.1节的三个流程级开关控制。')

NEW('不符合准入条件时的界面表现：签约入口仅显示 Manual Acceptance 一个选项，不显示 E-Acceptance 选项（而非显示后置灰）。符合准入条件时同时显示两个选项，默认选中 E-Acceptance。')

H(2, '2.4 生物识别对象与准入矩阵')
NEW('本节为新增，对应原始 FRS HP047 第5章的两张可配置矩阵。')

H(3, '2.4.1 生物识别对象角色')
P('生物识别的验证对象包含三类角色，不止主申请人与担保人：')
T(['角色', '英文', '说明'], [
    ['主申请人', 'Primary Applicant', '贷款借款人本人'],
    ['担保人', 'Guarantor', '为贷款提供担保的第三方'],
    ['车主/物权人', 'Owner', '车辆登记所有人。原始FRS矩阵中，企业类客群（Commercial Bank / Company / Partnership 等）的 Owner 角色需要做生物识别 = YES'],
], widths=[2.5, 4.0, 9.0])

TODO('P-23', 'Owner 角色在本次改造中是否保留？原始 FRS 中企业类客群的 Owner 需做生物识别，但本文档 2.2 节将 Non-Individual 统一归为「无 Biometric」，二者冲突。需确认车辆登记人与借款人不一致时的处理方式。')

H(3, '2.4.2 准入矩阵（BIOMETRIC_CRITERIA_CHECK）')
P('决定某人是否「需要」做生物识别。维度：客户类型 × BNM Group Code × 与申请的关系 × 证件类型。原始 FRS 中该矩阵为可配置表，本次改造需保留可配置能力。')
T(['客户类型', 'Primary Applicant', 'Guarantor', 'Owner'], [
    ['Individual（ID Card Blue/Red）', 'YES', 'YES', 'YES'],
    ['Corporate 各类（Bank / Company / Partnership 等）', 'N/A', 'N/A', 'YES'],
], widths=[6.5, 3.2, 2.8, 3.0])

H(3, '2.4.3 强制矩阵（BIOMETRIC_MANDATORY_CHECK）')
P('决定「未完成生物识别是否阻断流程」。命中 YES 时，未完成则工作流校验器拦截，案件无法推进（错误提示见第13章）。原始 FRS 中该矩阵同样可配置。')
T(['客户类型', 'Primary Applicant', 'Guarantor', 'Owner'], [
    ['Individual', 'YES', 'YES', 'YES'],
    ['Corporate', 'N/A', 'N/A', 'N/A'],
], widths=[6.5, 3.2, 2.8, 3.0])

P('补充规则：多个 Owner 的申请，系统只需校验其中一个 Owner（该逻辑不可配置）。')

H(3, '2.4.4 卖家（Direct）生物识别')
NEW('原始 FRS 中，生物识别搜索页列出的是「所有申请人及卖家（Direct）」，由开关 BIOMETRIC_CHECK_SELLER 控制，取值 I（个人卖家）/ C（企业卖家）/ OFF（关闭）。一旦开启，该卖家的生物识别即为强制完成项。')

TODO('P-24', '卖家（Seller Direct）生物识别功能本次是否保留？若保留，需确认开关默认值及适用的卖家类型。')

H(2, '2.5 外籍客户豁免与替代材料')
P('外籍客户（持 Passport）豁免生物识别，改由文件替代核验身份。')
TODO('P-25', '外籍客户替代材料的具体要求待确认：客户反馈提到可上传 Passport 或 Birth Certificate 替代，需明确①两者是二选一还是均需提供；②上传位置（Document Centre 或 File Attachment）；③是否需要经办人二次确认。')

doc.add_page_break()

# =====================================================================
# 3 Acceptance 状态模型
# =====================================================================
H(1, '3. Acceptance 状态模型')

H(2, '3.1 字段定义')
T(['字段', '字段名', '取值范围', '说明'], [
    ['Acceptance Mode', 'acceptance_mode', 'blank / Manual Acceptance / E-Acceptance', '标识签约模式：未确定 / 人工 / 电子'],
    ['Acceptance Status', 'acceptance_status', 'blank / Pending / Completed', '标识签约状态：未开始 / 进行中 / 已完成'],
    ['Failure Reason', 'failure_reason', 'blank / Biometric Failed / OTP Failed / RIB Unavailable / Not Registered', '失败原因标记，Mode=E-Acceptance 且 Status=Pending 时有效'],
], widths=[3.2, 3.0, 5.5, 3.8])

P('设计说明：Fail 不再作为独立状态，而是 E-Acceptance Mode + Pending Status 下的异常标记，系统通过 failure_reason 字段记录具体失败类型。')

H(2, '3.2 状态组合')
T(['业务场景', '状态组合', '说明'], [
    ['未开始', 'Mode=blank, Status=blank', '尚未进入签约步骤'],
    ['电子签进行中', 'Mode=E-Acceptance, Status=Pending, failure_reason=blank', '进入 E-Acceptance 流程，正常推进中'],
    ['电子签已失败', 'Mode=E-Acceptance, Status=Pending, failure_reason≠blank', 'E-Acceptance 过程中失败，待转 Manual'],
    ['人工签进行中', 'Mode=Manual Acceptance, Status=Pending', '进入 Manual Acceptance 流程'],
    ['签约完成', 'Mode=E-Acceptance / Manual Acceptance, Status=Completed', '签约完成'],
], widths=[3.0, 7.5, 5.0])

H(2, '3.3 列表页展示规则')
FIX('v1.0 的展示规则表未为失败场景保留展示行，导致「正在进行中」与「已失败待转手动」在列表页均显示为 E-Acceptance:Pending，运营与 CRA 无法分拣。本版增加失败态展示文本。')

T(['Acceptance Mode', 'Acceptance Status', 'Failure Reason', '列表页展示文本', '含义'], [
    ['blank', 'blank', '—', '（空）', '尚未进入签约步骤'],
    ['Manual Acceptance', 'Pending', '—', 'Manual: Pending', '人工签约进行中'],
    ['Manual Acceptance', 'Completed', '—', 'Manual: Completed', '人工签约已完成'],
    ['E-Acceptance', 'Pending', 'blank', 'E-Acceptance: Pending', '电子签约进行中'],
    ['E-Acceptance', 'Pending', '≠blank', 'E-Acceptance: Failed', '电子签失败，需转人工签'],
    ['E-Acceptance', 'Completed', '—', 'E-Acceptance: Completed', '电子签约已完成'],
], widths=[3.2, 2.6, 2.2, 4.0, 3.5])

P('说明：列表页不区分 E-Acceptance (Full) 与 Manual Acceptance (Hybrid)，Hybrid 在列表页统一显示为 Manual。Full/Hybrid 的区分在详情页和流程步骤中体现。')

TODO('P-26', '列表页失败态的展示文本需与客户确认：本版建议 "E-Acceptance: Failed"，另一方案是沿用原始 FRS 的独立状态值 "Fail"。需确认是否影响既有报表口径。')

H(2, '3.4 状态需覆盖的界面')
P('上述状态字段需在第10章列出的五个界面同步展示，详见 10.1 节。')

doc.add_page_break()

# =====================================================================
# 4 签约模式切换规则  ★NEW（0901评审）
# =====================================================================
H(1, '4. 签约模式切换规则')
NEW('本章为 2026-09-01 评审新增。此前切换规则散落在文档各处，本次统一汇总并前置。')

H(2, '4.1 前提条件')
P('只有当案件**同时具备 E-Acceptance 与 Manual Acceptance 两种资格**时，才存在「切换」这一动作。')
P('若案件仅满足 Manual 资格，签约入口**只显示 Manual Acceptance 一个选项**——不显示 E-Acceptance 卡片，也不存在置灰、高亮或默认选中的处理。')

H(2, '4.2 切换矩阵')
T(['#', '场景', '能否切换', '路径与限制'], [
    ['1', '签约尚未完成（任意中间状态）', '✅ 双向可切',
     'Manual → E-Acceptance、E-Acceptance → Manual 均允许。E-Acceptance 流程卡在任何一步（第1次生物识别 / RIB 阅读合同 / 第2次生物识别）都可切换。'],
    ['2', '客户已在 RIB 点击 Submit，但第2次验证未完成', '✅ 可切',
     '签署动作虽已发生，但只要第2次生物识别或 OTP 未成功，签约即未完成，仍可切至 Manual。'],
    ['3', 'Manual 已完成、尚未提交放款', '✅ 可切回 E',
     'Sales 在文件中心**删除**已上传的 CRA HP Agreement → 系统重置 Acceptance 状态回 Pending → 重新进入模式选择入口。'],
    ['4', '已提交放款', '❌ 不可切', '文件已不可删除，切换路径自然关闭。'],
    ['5', 'E-Acceptance 失败达上限被强制转 Manual', '❌ 本轮锁定',
     '本轮 Acceptance Cycle 内不可再选择 E-Acceptance。失败次数须由后端持久化记录，刷新页面不可绕过。'],
    ['6', '签约已 Completed', '❌ 不可切', '签约完成即锁定，不允许变更签约方式。'],
], widths=[1.0, 4.0, 2.4, 8.1])

H(2, '4.3 与原 LOAD$ 做法的差异')
P('原 LOAD$ 的 Manual → E-Acceptance 转换方式是：上传一份 CRA HP Agreement (E-Acceptance) 文件，以此强制翻转签约状态。')
FIX('本次改造**不采用**该方式。原因：新系统中 E-Acceptance 文书为系统自动生成并自动归档，不需要人工上传。改为「删除 Manual 上传件即重置状态」。')

H(2, '4.4 首次签约类型对后续 CILT 的影响')
T(['首次签约类型', 'CILT 后是否需重新签约', 'Acceptance 状态', '后续可否更换模式'], [
    ['E-Acceptance', '需要（当 CILT 改动涉及签署文书字段时）', '更新为 Pending', '✅ 可更换为 Manual Acceptance'],
    ['Manual Acceptance', '**不需要**', '保持 Completed 不变', '❌ 不可更换为 E-Acceptance'],
], widths=[3.2, 5.0, 3.3, 4.0])

P('判断依据：CILT 所修改的字段是否涉及需客户签署的文书。新车为 3 份（HP Agreement、Second Schedule Part 1、Product Disclosure Sheet），二手/翻新车为 4 份（增加 Appendix 4）。若改动未涉及上述文书字段，则 Acceptance 状态不变，始终保持 Completed。')

H(2, '4.5 强制切换的落库要求')
P('所有「由系统强制从 E-Acceptance 切换至 Manual Acceptance」的场景，其触发原因与失败计数**必须由后端持久化**，不得仅依赖前端状态控制。用户刷新页面或重新进入签约入口时，不得绕过该限制重新选择 E-Acceptance。')

doc.add_page_break()

# =====================================================================
# 5 E-Acceptance (Full) 流程
# =====================================================================
H(1, '5. E-Acceptance (Full) 流程')

H(2, '5.1 适用条件')
P('E-Acceptance (Full) 仅在第2.3节四项准入条件全部满足时适用。流程特征：Biometric Verification ✅ + E-signing ✅，全程线上完成。')

H(2, '5.2 流程概览（7步）')
T(['步骤', '名称', '执行方', '关键动作'], [
    ['Step 1', '系统进入签约流程', '系统', '审批通过后自动判断客群分流'],
    ['Step 2', '签约入口显示', 'Sales', '选择「Proceed with E-Acceptance」，生成待签文件'],
    ['Step 3', '第1次 Biometric（Identity Verification）', 'Sales + 客户', '验证客户身份，成功后解锁 RIB'],
    ['Step 4', 'RIB 网银阅读确认', '客户', '逐份阅读并确认4份文件后 Submit'],
    ['Step 5', '第2次 Biometric（E-Acceptance Confirmation）', 'Sales + 客户', '确认客户签约意愿'],
    ['Step 6', 'OTP 短信验证（兜底）', '客户', '仅在第2次 Biometric 失败时触发'],
    ['Step 7', '签约完成', '系统', '自动归档、状态流转、客户收到通知'],
], widths=[1.8, 5.2, 3.0, 5.5])

H(2, '5.3 各步骤详解')

H(3, 'Step 1：系统进入签约流程')
T(['字段', '说明'], [
    ['触发时机', '贷款审批通过后，主单状态 Approved，进入签约流程'],
    ['执行逻辑', '根据客户类型 / 国籍 / 担保人 / 渠道判断，确定可选的签约模式'],
    ['状态变更', 'Mode=blank → 依判定结果预置'],
], widths=[3.0, 12.5])

H(3, 'Step 2：签约入口显示')
P('案件通过审批后，Sales 的签约入口打开。对于符合 E-Acceptance 资格的案件，展示电子签与人工签两种方式，默认选中电子签；对于不符合资格的案件，仅显示 Manual Acceptance 一个选项。')
P('同时生成待签约文件：HP Agreement、Second Schedule Part 1、Product Disclosure Sheet；条件性附加 Appendix 4（仅适用于二手/翻新车）。')

H(3, 'Step 3：第1次 Biometric（Identity Verification）')
T(['字段', '说明'], [
    ['执行方', 'Sales 操作，客户按指纹'],
    ['系统', 'CrediOS + 生物识别设备'],
    ['验证对象', 'Primary Applicant（主申请人）'],
    ['结果状态', 'Matched / Unmatched / Error / Not Done'],
    ['成功后续', '系统生成 RIB 门户跳转链接，客户可登录 RIB 完成 E-signing'],
    ['失败后续', '默认允许重试3次；达上限后强制转 Manual Acceptance，且本轮不可再选 E-Acceptance（见第4章）'],
    ['次数可配置性', '「失败几次后转 Manual」由 BIOMETRIC_ATTEMPT_COUNTER 配置，默认 3 次，系统上限 999。OTP 失败转 Manual 的次数同样可配置'],
    ['超时', '无超时限制，Sales 可随时选择转 Manual'],
], widths=[3.0, 12.5])

FIX('v1.0 中重试次数在四处分别写作 3 次与 999 次且同句矛盾，本版统一为「默认3次，可配置，系统上限999」。')

P('【前端页面设计】生物识别详情页（Biometric Scanning Details Screen）分为两部分：主信息页（展示与记录）和实操扫描区。')

P('主信息页左侧卡片展示字段：', bold=True)
T(['字段', '说明'], [
    ['App. Ref. No.', '申请编号'],
    ['Name', '客户姓名'],
    ['ID No.', '身份证号'],
    ['Verification Type', '验证类型（Identity Verification / E-Acceptance Confirmation / Manual Acceptance Verification）'],
    ['Staff ID', '当前操作扫描的柜员工号'],
    ['Date & Time Request', '发起验证请求的时间（多次触发只取最新一条）'],
    ['Date & Time Response', '收到验证结果的时间（多次触发只取最新一条）'],
    ['Biometric Status', 'Matched / Unmatched / Error / Not Done'],
    ['Triggering Point', '触发时所在的工作流步骤（若跨多步骤触发，全部列出，取最新记录）'],
    ['Triggering Counter', '累计触发次数（含 Error）'],
    ['Remarks', '自由文本备注；Status 为 Unmatched 或 Error 时必填，否则校验器拦截'],
], widths=[4.2, 11.3])

NEW('v1.0 的字段表缺 App. Ref. No.、Name、ID No.、Verification Type 四项（客户反馈第7条明确要求），本版补齐。')

P('实操扫描区（Scanning Operation Panel）：点击「Start Biometric Scan」后进入，包含以下元素：', bold=True)
T(['元素', '说明'], [
    ['Read MyKad', '触发读卡器读取客户身份证（MyKad）'],
    ['Restart Service', '设备异常时重启生物识别服务。仅在 Error 状态下出现，非常驻按钮'],
    ['Back', '保存当前结果并返回主信息页（注意：不是取消操作）'],
    ['Status Textbox', '动态状态提示框，实时显示硬件状态'],
], widths=[3.6, 11.9])

P('Status Textbox 文案（沿用原始 FRS 原词）：', bold=True)
T(['场景', '文案'], [
    ['就绪', 'Reader is ready'],
    ['引导按指纹', 'Please place your thumb on the scanner when light is On'],
    ['验证通过', 'Verified OK. Please Remove Card. / Thumbprint matched.'],
    ['未检测到卡', 'No MyKad'],
    ['操作超时', 'Error - Time Out'],
    ['设备未连接/驱动异常', 'Error fetching listing'],
], widths=[4.5, 11.0])

P('主信息页底部操作按钮：Start Biometric Scan（进入实操扫描区）、Report（生成并下载 Biometric Report）、Save（保存备注）、Back（返回上一页）。')

P('流程联动：第1次指纹匹配成功（Matched）后，系统解锁并跳转至 RIB 门户，客户开始阅读合同文件。若不匹配或失败，RIB 门户不会弹出。')

H(3, 'Step 4：RIB 网银阅读确认')
T(['字段', '说明'], [
    ['执行方', '客户（Primary Applicant）'],
    ['系统', 'RIB（Retail Internet Banking）'],
    ['操作', '客户登录 RIB → 逐份查看协议文件 → 确认签署'],
    ['签署内容', 'HP Agreement + Second Schedule Part 1 + Product Disclosure Sheet + Appendix 4（二手/翻新车适用）'],
    ['回传机制', 'RIB 签署完成后，回传 Acceptance Flag 至 CrediOS'],
], widths=[3.0, 12.5])

P('【前端设计规范】RIB 门户合同审阅页（Contract Review Page）：', bold=True)
NUM('待审阅合同列表：页面清晰列出需客户逐一阅读的文件（共3–4份，Appendix 4 仅二手/翻新车显示）')
NUM('客户必须逐一点击进入每份合同，滚动到底部后勾选「I hereby declare that I have read and understood」复选框，再点击「Confirm」确认')
NUM('每份合同确认后，返回主列表时该条目旁显示绿色勾号（Green Tick ✓），未阅读的条目保持空白；已确认的文件仍可点击重新查看，确认记录不受影响')
NUM('主列表底部的 Submit 按钮默认置灰，只有所有文件都获得绿勾后才解锁')
NUM('客户点击 Submit 后，页面显示提示「Your bank officer will now scan your fingerprint to complete the acceptance」，RIB 页面进入等待状态')

H(3, 'Step 5：第2次 Biometric（E-Acceptance Confirmation）')
T(['字段', '说明'], [
    ['触发条件', 'CrediOS 收到 RIB 回传的 Acceptance Flag'],
    ['执行方', 'Sales 操作，客户按指纹'],
    ['验证对象', 'Primary Applicant（主申请人）'],
    ['超时', '15分钟（从触发第2次 Biometric 开始计时）'],
    ['成功后续', 'RIB 门户无缝跳转到「Thank You for accepting the loan/financing documents」成功页，允许客户下载已签合同；Acceptance Status → Completed'],
    ['失败后续', '第2次 Biometric 失败或超时 → 转入 OTP 验证（Step 6）'],
], widths=[3.0, 12.5])

P('操作路径：销售回到 CrediOS，选择「E-Acceptance Confirmation」并点击开始扫描，客户进行第二次按指纹。')

H(3, 'Step 6：OTP 短信验证（兜底）')
FIX('2026-09-01 评审确认：OTP 全流程由 RIB 侧负责，CrediOS 不实现 OTP 的发送、输入与校验。此前文档与原型中「CrediOS 弹出 OTP 输入框」的描述有误，本版更正。')

T(['环节', '归属系统'], [
    ['触发 OTP 发送页面', 'RIB'],
    ['客户点击 Receive OTP', 'RIB'],
    ['OTP 短信下发', 'RIB'],
    ['客户输入验证码', 'RIB 页面'],
    ['验证码校验', 'RIB'],
    ['Resend OTP', 'RIB'],
], widths=[6.0, 9.5])

T(['字段', '说明'], [
    ['触发条件', '仅在第2次 Biometric 匹配失败或超时时触发'],
    ['限制', '首次身份验证（Identity Verification）失败不能使用 OTP 兜底'],
    ['CrediOS 职责', '通过接口通知 RIB「第2次生物识别失败」（参数 iqType = second bmt failed），随后等待 RIB 返回结果'],
    ['CrediOS 界面表现', '显示「OTP 验证进行中（RIB 侧）」等待态，说明客户需在 RIB 页面完成验证；提供 Refresh 按钮供操作员刷新结果；不显示 OTP 输入框'],
    ['OTP 参数', '6位数字验证码，10分钟有效期（由 RIB 侧控制）'],
    ['成功记录', 'OTP 验证成功记录终生有效'],
    ['成功后续', 'RIB 回传验证成功 → Acceptance Status → Completed'],
    ['失败后续', 'OTP 多次失败 → 转 Manual Acceptance'],
], widths=[3.4, 12.1])

TODO('P-33', '🔴 **阻塞项**：目前尚未找到 RIB 向 CrediOS 回传「OTP 验证结果」的接口。若该接口缺失，CrediOS 将无法感知客户是否完成验证，流程会停滞在等待态。已向客户科技提出，待回复。')

P('【RIB 侧前端表现，供参考】指纹识别失败后，RIB 页面弹出错误提示框（带手机插画），显示文案：「很抱歉，我们无法获取您的指纹。请点击下方『获取 OTP』并输入发送到您手机的 6 位验证码」。客户点击「Receive OTP」按钮，输入手机收到的 6 位数字验证码。')

H(3, 'Step 7：签约完成')
T(['字段', '说明'], [
    ['触发条件', '第2次 Biometric 成功 或 OTP 验证成功'],
    ['自动操作1', '系统调用 attachLetter 函数，将签署完成的文件盖章后自动附加到 File Attachment'],
    ['自动操作2', '生成 E-Acceptance Print Log PDF'],
    ['自动操作3', 'E-Acceptance 文件自动归档至 EDMS'],
    ['状态变更', 'Mode=E-Acceptance, Status=Completed'],
    ['主单状态', 'Approved → Accepted'],
    ['后续步骤', 'Sales 点击「Proceed with E-Acceptance」→ 案件流转至 Pending Funding Document → 补齐剩余放款文件 → 点击「Route to CRA checker」进入放款流程'],
], widths=[3.0, 12.5])

NEW('v1.0 的后续步骤仅写「proceed to CRA」，本版补充主单状态从 Approved 变为 Accepted 的状态跃迁，以及放款推进的完整三步。若必填放款文件未挂齐，系统在 Route to CRA checker 时报错拦截。')

P('自动挂载文件清单（5类）：', bold=True)
T(['文档', '说明'], [
    ['HP Agreement', '汽车贷款协议（电子签版）'],
    ['Second Schedule Part 1', '第二附表第一部分'],
    ['Product Disclosure Sheet', '产品披露表'],
    ['Customer Biometric Scanning Result', '客户生物识别结果单。电子签场景下不需要柜员单独提交，系统自动挂载'],
    ['Appendix 4', '附录4，仅二手/翻新车适用'],
], widths=[5.5, 10.0])

H(2, '5.4 异常分支')
P('详见第20章异常分支汇总（A-1 至 F-3）。')

H(2, '5.5 通用交互规则')
NEW('2026-09-01 评审新增。')

H(3, '5.5.1 浏览器刷新与断点')
T(['场景', '刷新后的表现'], [
    ['第1次生物识别进行中', '停留在第1次生物识别步骤'],
    ['已跳转 RIB、等待客户签署', '停留在等待 RIB 回执步骤'],
    ['第2次生物识别进行中', '停留在准备第2次生物识别步骤'],
    ['CILT — 已完成 Step 1（已同步并发短信）', '停留在 Step 2，**不得回退**至 Step 1 重复发送'],
], widths=[6.0, 9.5])

P('要求：上述业务步骤须由后端持久化记录，Sales 刷新 CrediOS 浏览器页面或重新进入案件时，均回到当前实际业务步骤，不得回退或跳过。', bold=True)

H(3, '5.5.2 Back 按钮规则')
T(['所处阶段', '底部 Back', '说明'], [
    ['生物识别、发送 RIB 等流程步骤中', '✅ 保留', '供 Sales 返回上一步'],
    ['CILT / Amendment 阶段', '❌ 不显示', '上一轮签约已完成，底部不需要操作按钮'],
    ['签约完成态', '仅保留一个返回签约入口的 Back', '其余底部按钮全部移除'],
], widths=[5.5, 3.5, 6.5])

P('生物识别详情页的 Back 按钮位于左侧字段区，用途为「保存当前结果并返回」，不是取消操作。')

H(3, '5.5.3 Switch to Manual Acceptance 按钮')
P('E-Acceptance 流程中，该按钮**常驻于页面右上角**，在签约完成前的任意步骤均可点击。各失败态（如重试区域）**不再重复放置**该按钮。')

doc.add_page_break()

# =====================================================================
# 5 Manual Acceptance (Hybrid)
# =====================================================================
H(1, '6. Manual Acceptance (Hybrid) 流程')

FIX('本章所有状态变更的 Mode 取值，v1.0 中误写为 E-Acceptance（共5张表约12处），本版统一修正为 Manual Acceptance。')

H(2, '6.1 适用条件')
P('流程特征：Biometric Verification + Paper Signing，认证在线、签署线下。适用客群：')
T(['客群编号', '渠道', '担保人', '国籍'], [
    ['I-3', 'Dealer', '有', '任意'],
    ['I-4', 'Direct', '无', '本地'],
    ['I-6', 'Direct', '有', '任意'],
], widths=[3.0, 3.5, 3.5, 5.5])

H(2, '6.2 与 Full 流程的差异点')
T(['差异点', 'E-Acceptance (Full)', 'Manual Acceptance (Hybrid)'], [
    ['Signing 方式', 'E-signing（RIB网银在线签署）', 'Paper Signing（纸质签字）'],
    ['第2次验证', 'Biometric 或 OTP', '不需要（Paper Signing 本身即确认）'],
    ['RIB 签署', '需要', '不需要'],
    ['OTP 验证', '需要（作为备选）', '不需要'],
    ['担保人 Biometric', '不适用（无担保人）', '需要（I-3 / I-6 客群）'],
    ['文件上传', '系统自动附加', 'Sales 需手动上传签署后的纸质文件'],
], widths=[3.5, 6.0, 6.0])

H(2, '6.3 各步骤详解')

H(3, 'Step 1：签约入口显示，选择人工签')
T(['字段', '说明'], [
    ['执行方', 'Sales'],
    ['系统', 'CrediOS'],
    ['操作', 'Sales 选择 Manual Acceptance'],
    ['状态变更', 'Mode=Manual Acceptance, Status=Pending'],
], widths=[3.0, 12.5])

H(3, 'Step 2：第1次 Biometric（Identity Verification）')
T(['字段', '说明'], [
    ['验证对象', '主申请人 + 担保人（I-3 / I-6 客群）；或仅主申请人（I-4 客群）'],
    ['执行顺序', '主申请人先完成，担保人后完成，依次进行'],
    ['成功后续', '系统提示 Sales 准备纸质协议文件'],
    ['失败后续', '同 Full 流程异常 A'],
    ['状态变更', 'Mode=Manual Acceptance, Status=Pending'],
], widths=[3.0, 12.5])

H(3, 'Step 3：Paper Signing（纸质签字）')
T(['字段', '说明'], [
    ['执行方', 'Sales 打印协议文件，客户在纸质协议上亲笔签字'],
    ['签署对象', '主申请人 + 担保人（如适用）'],
    ['说明', 'Hybrid 不需要第2次 Biometric 和 OTP，Paper Signing 本身即代表客户的签约确认'],
], widths=[3.0, 12.5])

H(3, 'Step 4：上传签署完成文件')
P('签署完成后，Sales 将签署后的纸质文件扫描上传至 CrediOS File Attachment。上传时的文档类型选择会影响案件签约类型打标，详见第10.4节。')

H(3, 'Step 5：签约完成')
T(['字段', '说明'], [
    ['触发条件', 'Sales 上传签署后的纸质文件'],
    ['自动操作1', '生成 Print Log PDF'],
    ['自动操作2', '文件自动归档至 Attachment'],
    ['状态变更', 'Mode=Manual Acceptance, Status=Completed'],
    ['后续步骤', 'Sales Pending Fund Document 步骤 → 放款（Funding）'],
], widths=[3.0, 12.5])

H(2, '6.4 生物识别跳过与补件规则')
T(['类型', '对象', '系统内生物识别', '放款前补件'], [
    ['强制 + 自动归档', '本地个人主申请人（无担保人、经销商购车）', '必须完成', '自动附档，无需补件'],
    ['跳过 + 必须补件', '所有担保人；转为人工签的主申请人', '可跳过系统硬件验证', '强制上传纸质指纹凭证至 Document Centre'],
    ['完全豁免 N/A', '外籍人士；企业/公司客户', '免', '免'],
], widths=[3.2, 5.3, 3.2, 3.8])

H(2, '6.5 异常分支')
P('详见第19.2节（H-1 至 H-6）。')

doc.add_page_break()

# =====================================================================
# 6 Manual Acceptance
# =====================================================================
H(1, '7. Manual Acceptance 流程')

H(2, '7.1 适用条件')
P('流程特征：无在线 Biometric 验证，纯纸质签字。适用客群：')
T(['客群编号', '条件'], [
    ['I-2', 'Individual + Dealer + 无担保人 + 外籍（Biometric 豁免）'],
    ['I-5', 'Individual + Direct + 无担保人 + 外籍（Biometric 豁免）'],
    ['Non-Individual', '非个人申请'],
], widths=[3.5, 12.0])

H(2, '7.2 流程详细说明')
T(['字段', '说明'], [
    ['触发场景', '① 系统分流结果为 Manual；② Sales 主动选择 Manual；③ E-Acceptance 流程中 Biometric 失败超限；④ 客户未注册 RIB；⑤ 客户要求纸质签约'],
    ['流程步骤', 'Sales 打印协议文件 → 客户线下签署 → Sales 上传签署文件'],
    ['Biometric', 'Manual Acceptance Verification 类型，在人工接受过程中仍需进行身份验证（I-2 / I-5 客群豁免）'],
    ['状态记录', 'Mode=Manual Acceptance'],
], widths=[3.0, 12.5])

H(2, '7.3 文件上传要求（6类材料）')
T(['文档类型', '说明', '是否必须'], [
    ['CRA HP Agreement (Manual Acceptance)', '人工签约协议', '必须'],
    ['客户身份证明复印件', '客户 ID 副本', '必须'],
    ['签署协议原件扫描件', '纸质协议扫描件', '必须'],
    ['担保人身份证明（如适用）', '担保人 ID 副本', '条件必须'],
    ['Appendix 4（如适用）', '附录4，仅二手/翻新车', '条件必须'],
    ['其他补充材料', '额外支持文件', '可选'],
], widths=[6.0, 6.0, 3.5])

H(2, '7.4 Manual → E-Acceptance 转换')
T(['字段', '说明'], [
    ['触发条件', '符合电子签条件的案件被误操作走成 Manual 流程时，支持纠错转回 E-Acceptance'],
    ['操作步骤', '上传文档类型为 CRA HP Agreement (E-Acceptance) 的文件至系统'],
    ['系统行为', '弹窗二次确认 → 确认后系统自动抓取最新版 HP Agreement 挂载（Browse 按钮禁用，无需实际上传文件）'],
    ['状态变更', 'Mode 从 Manual Acceptance 变更为 E-Acceptance'],
    ['限制条件', '仅在放款前可转换；放款后不可逆'],
], widths=[3.0, 12.5])

P('完整的四场景签约类型打标机制详见第10.4节。')

H(2, '7.5 异常分支')
P('详见第19.3节（M-1 至 M-4）。')

doc.add_page_break()

# =====================================================================
# 7 Amendment
# =====================================================================
H(1, '8. Amendment 修订流程')

H(2, '8.1 适用场景与触发规则')
P('当贷款条件发生变更（如利率调整、期限变更、月供调整等）时，已签署的协议需要修订并重新签署。Amendment 流程适用于所有已进入 Acceptance 流程的申请（不论 Full / Hybrid / Manual）。')

T(['场景', '处理方式'], [
    ['首次签约为 E-Acceptance，CILT 发生在签约完成之后', '按 Amendment 修订流程处理，需客户重新确认（本章主流程）'],
    ['首次签约为 Manual Acceptance', '**不需要重新签约**，Acceptance 状态保持 Completed，直接推进放款（见 8.7 节）'],
    ['CILT 发生在客户尚未做任何签约动作之前', '案件退回签约步骤，直接当作全新的「初次签约」案件处理，不算修订'],
    ['CILT 改动未涉及签署文书字段', 'Acceptance 状态不变，无需重新签约'],
], widths=[6.5, 9.0])

P('前置条件：Amendment 请求仅在第2次 Biometric（E-Acceptance Confirmation）完成后才适用。')

H(2, '8.2 流程概览')
T(['步骤', '执行方', '动作'], [
    ['1', 'Sales', '在 CrediOS 中修改相关字段（CILT）'],
    ['2', '系统', '自动追踪变更字段（比对上一次已接受版本）'],
    ['3', '系统', '重新生成协议文件，变更字段高亮标注'],
    ['4', 'Sales', '**Step 1**：点击「Sync & Notify Customer」——同步 e-Tracker 并通知客户（见 8.3 节）'],
    ['5', '客户', '通过 e-Tracker 远程审阅并确认（见 8.4 节）'],
    ['6', 'Sales', '**Step 2**：点击「Refresh」刷新，获取客户确认结果'],
    ['7', '系统', '在 Acceptance Log 盖一条 Amendment Acceptance 记录'],
], widths=[1.5, 2.5, 11.5])

H(2, '8.3 Sales 侧操作：两步流程')
NEW('2026-09-01 评审决定：原「① 同步 e-Tracker → ② 发送短信 → ③ 刷新状态」三个独立按钮，合并为两步。')

H(3, 'Step 1：Sync & Notify Customer（同步并通知）')
P('单一按钮完成两件事，由后台按时序执行：')
T(['顺序', '动作', '失败处理'], [
    ['①', '同步修订信息至 e-Tracker', '同步失败 → **不发送短信**，界面**停留在 Step 1**，Acceptance 状态不刷新，接口返回错误提示「同步失败，请稍后重试」'],
    ['②', '通过 Infobip 网关发送短信通知客户', '短信发送失败 → 界面推进至 Step 2，短信状态标记为「发送失败」，由 Resend 按钮补发'],
], widths=[1.5, 5.5, 8.5])

P('记录与展示规则：', bold=True)
BULLET('**e-Tracker 同步结果不记录日志、不在界面展示**（评审结论：避免日志冗余）')
BULLET('**仅记录短信发送状态**（成功 / 失败），并在界面展示')

P('客户信息（姓名、手机号、CILT 编号）在本步骤顶部展示，供 Sales 核对后再发送。')

H(3, 'Step 2：Awaiting Customer Confirmation（等待客户确认）')
T(['元素', '说明'], [
    ['短信状态', '展示「已发送 + 发送时间」或「发送失败」'],
    ['Resend 按钮', '短信发送失败时用于补发。点击后显示 loading，接口返回成功后**自动刷新本屏全部展示内容**'],
    ['Refresh 按钮', '刷新**本屏需要展示的全部内容，包含短信发送状态**，并获取客户是否已确认'],
    ['自动刷新', '每次进入该界面时可自动执行一次刷新，不强制要求 Sales 手动点击'],
], widths=[3.2, 12.3])

P('断点要求：Sales 完成 Step 1 后离开页面，再次进入时应停留在 Step 2，不得回退至 Step 1 重复发送。该断点须由后端持久化记录。', bold=True)

FIX('客户在 e-Tracker 完成确认后，CrediOS **不会自动推进状态**——需 Sales 点击 Refresh 主动获取结果。原型中的 e-Tracker 客户视图弹窗仅为演示用途，真实系统在该位置只有一个 Refresh 按钮。')

H(2, '8.4 e-Tracker 客户端远程重签交互（仅 E-Acceptance 客群）')

H(3, '客户操作四步')
NUM('收到短信，点击其中的 e-Tracker 链接')
NUM('登录 e-Tracker 后，在该笔贷款进度条下方看到提示文字，并出现醒目的红色「Review Now」按钮')
NUM('点击进入合同页面，所有变动字段以**黄色背景高亮**（Yellow Highlight）展示')
NUM('勾选底部声明「I hereby declare that I have read and understood the final loan documents and the terms and conditions」，点击红色「Confirm」按钮完成远程重签')

P('适用范围：该流程仅适用于 E-Acceptance 客群的 CILT Amendment；Manual Acceptance 客群不需要重新签署。')

H(2, '8.5 变更追踪字段完整清单')
NEW('本节为新增。v1.0 将此项列为待确认 P-07，实际原始 FRS HP075 第7.1节已按四个模板逐字段定义，共46个字段，现予补齐并结案。')

H(3, '8.5.1 HP Agreement Part 1（13个字段）')
T(['#', '字段', '#', '字段'], [
    ['1', 'Name', '8', 'Total Amount'],
    ['2', 'Address', '9', 'Balance originally payable under this agreement'],
    ['3', 'Description of Goods', '10', 'Annual Percentage Rate of Term Charges'],
    ['4', 'New/Second-hand', '11', 'Duration of payment by installment (months/years)'],
    ['5', 'Goods to be kept at', '12', 'Number of Installment'],
    ['6', 'Financial Amount', '13', 'Final Installment'],
    ['7', 'Interest Rate', '', ''],
], widths=[1.0, 6.5, 1.0, 7.0])

H(3, '8.5.2 HP Second Schedule Part 1（21个字段）')
T(['#', '字段', '#', '字段'], [
    ['1', 'Full Name', '12', 'Rate per annum'],
    ['2', 'Short description of goods', '13', 'Total amount of term charges / mark-up'],
    ['3', 'Registration number', '14', 'Balance originally payable under the agreement'],
    ['4', 'New/Second Hand', '15', 'Annual Percentage Rate'],
    ['5', 'Address where goods will be kept', '16', 'Hire Purchase Price'],
    ['6', 'Cash price of goods', '17', 'Difference between cash price and total payable'],
    ['7', 'Deposit', '18', 'Duration of Payment of instalment'],
    ['8', 'Cash price less deposit', '19', 'Number of Installment'],
    ['9', 'Freight charges', '20', 'Amount of each Installment'],
    ['10', 'Vehicle registration fee', '21', 'Final Installment'],
    ['11', 'Insurance / Takaful', '', ''],
], widths=[1.0, 6.5, 1.0, 7.0])

H(3, '8.5.3 Appendix 4 Customer/Hirer Indemnity Letter（5个字段）')
T(['#', '字段'], [
    ['1', 'Dealer/Vendor'],
    ['2', 'Vehicle Model'],
    ['3', 'Engine No.'],
    ['4', 'Chassis No.'],
    ['5', 'Full Name'],
], widths=[1.5, 14.0])

H(3, '8.5.4 Product Disclosure Sheet（7个字段）')
T(['#', '字段'], [
    ['1', 'Total amount financed'],
    ['2', 'Tenure'],
    ['3', 'Base Lending Rate (BLR)'],
    ['4', 'Interest Rate'],
    ['5', 'Annual Percentage Rate'],
    ['6', 'Your monthly instalment'],
    ['7', 'Total repayment amount'],
], widths=[1.5, 14.0])

H(2, '8.6 文件重新生成与高亮机制')
T(['规则', '说明'], [
    ['变更检测', '系统按7.4节字段清单，将新值与上一次已接受版本逐字段比对'],
    ['选择性推送', '只有检测到字段变化的文件才会重新生成并推送给客户二次签署；未变化的文件不重复要求客户确认'],
    ['多轮区分', '支持多轮修订，按修订轮次（amendment index）使用不同字体颜色高亮，并标记累计变更'],
    ['防覆盖', '修订后的文件作为新版本独立存储，原始签署文件永久保留'],
    ['文件附加', '修订后的签署文件自动附加到 File Attachment'],
], widths=[3.0, 12.5])

NEW('「选择性推送」与「多轮颜色区分」两条规则为本次补充，v1.0 仅写了「重新生成协议文件」，未区分哪些文件需要客户重签。')

TODO('P-08', '修订协议中变更字段高亮的具体样式规范待确认：包括黄色高亮的色值、多轮修订各轮次使用的字体颜色序列、以及累计变更的标记方式。')

H(2, '8.7 首次签约为 Manual Acceptance 时的 CILT 处理')
NEW('2026-09-01 评审新增，依据客户既有需求单（浩哥确认）。')

T(['项目', '规则'], [
    ['是否需要重新签约', '**不需要**。客户无需再次签署即可提交放款'],
    ['Acceptance 状态', '保持 **Completed** 不变，不回置为 Pending'],
    ['是否发送 e-Tracker 通知', '不发送。Manual 客群无 e-Tracker 重签流程'],
    ['后续可否改为 E-Acceptance', '**不可以**。首次为 Manual 的案件，后续各轮均维持 Manual'],
    ['文书处理', '系统仍重新生成更新后的文书并归档；各版本保留在 Document Library'],
    ['界面表现', 'CILT 界面显示「无需重新签约」说明，不出现同步/发送短信操作区'],
], widths=[4.0, 11.5])

H(2, '8.8 短信模板')
T(['配置项', '值'], [
    ['菜单', 'Parameter > Origins > SMS/Email Message Setup'],
    ['Product Type', 'HP'],
    ['Language', 'English'],
    ['Reason Type', 'Other Notification Reason'],
    ['Code', 'EAA'],
    ['Description', 'E-Acceptance Amendment Notification'],
    ['Gateway', 'Infobip'],
    ['Recipient', 'Primary Customer'],
    ['SMS Content', 'RM0.00 HLB: Dear customer, please review and accept the changes to your loan/financing documents via the link below. www.hlb.com.my/etracker'],
], widths=[3.5, 12.0])

doc.add_page_break()

# =====================================================================
# 8 身份验证机制
# =====================================================================
H(1, '9. 身份验证机制')

H(2, '9.1 生物识别模块定位与入口')
NEW('本节为新增。原始 FRS HP047 将生物识别设计为 LOAD$ 内的独立可复用模块，拥有独立菜单、搜索页与详情页，供 E-Acceptance 与 Manual Acceptance 共同调用，而非嵌在签约流程内的子功能。')

T(['界面', '说明'], [
    ['独立菜单入口', '生物识别为独立菜单项，适用用户权限矩阵控制可见性'],
    ['Biometric 搜索页', '列出所有需要执行生物识别的申请人及卖家（Direct）。支持按 App. Ref. No. 或 客户ID / 卖家ID 检索，不支持全量查询'],
    ['Biometric 详情页', '单个申请人的验证详情与操作区，详见 4.3 节 Step 3'],
    ['应用内 Biometric Tab', '在申请详情页内的标签页，支持多申请人记录切换'],
], widths=[4.0, 11.5])

TODO('P-27', '生物识别独立菜单的具体位置待确认。客户反馈提出应放在 e-Hakmilik 之上，需确认最终菜单层级与排序。')

H(3, '9.1.1 Biometric 搜索页字段')
T(['字段', '可排序', '说明'], [
    ['App. Ref. No.', 'Y', '申请编号，默认升序'],
    ['ID No.', 'Y', '按申请人 / 卖家 ID。支持用 ID1 或 ID2 检索，但结果始终显示 ID1'],
    ['Customer Name', 'Y', '客户姓名'],
    ['Relationship', 'Y', '申请人类型（Primary Applicant / Guarantor / Owner）或 Seller（Direct 卖家）'],
    ['Biometric Status', 'Y', 'Matched / Unmatched / Error（含描述）/ Not Done'],
    ['Triggered Counter', 'Y', '生物识别尝试次数（含 Error）'],
    ['Action', 'N', '点击进入 Biometric 详情页'],
], widths=[3.5, 2.0, 10.0])

H(3, '9.1.2 应用内 Biometric Tab')
NEW('原始 FRS HP047 4.4 定义的独立结构，v1.0 未覆盖。')
T(['区域', '字段/元素'], [
    ['记录选择区', 'Select（单选，默认选中第一条）、Applicant Type、ID No.、Name'],
    ['扫描详情区', 'Staff ID、Date & Time Request、Date & Time Response、Biometric Status、Triggering Point、Triggering Counter'],
    ['操作区', 'Remarks（可编辑）、Report（下载报表）、Save（保存备注）'],
], widths=[3.5, 12.0])

P('权限说明：Save 按钮的权限仅开放给处于签约步骤的 Sales，供其补填 Remarks。')

H(2, '9.2 验证类型')
T(['类型', '英文名称', '触发阶段', '目的', '适用流程'], [
    ['身份验证', 'Identity Verification', '第1次 Biometric', '验证客户身份', 'E-Acceptance (Full) / Hybrid'],
    ['签约确认', 'E-Acceptance Confirmation', '第2次 Biometric', '确认客户签约意愿', 'E-Acceptance (Full)'],
    ['人工接受验证', 'Manual Acceptance Verification', 'Manual Acceptance 流程', '人工接受中的身份验证', 'Manual Acceptance（非外籍）'],
], widths=[2.5, 4.0, 3.0, 3.0, 3.0])

P('第2次验证记录的生成条件（三者须同时满足）：① 该申请符合 E-Acceptance 准入；② 第1次验证结果为 Matched；③ 已收到 RIB 的 Acceptance Flag = YES。否则不生成第2次验证记录。')

H(2, '9.3 生物识别状态')
T(['状态', '含义', '后续操作'], [
    ['Matched', '指纹匹配成功', '继续后续流程'],
    ['Unmatched', '指纹不匹配', '核对是否申请错误或核实 MyKad 真伪；允许重试（受 Attempt Counter 限制）；必须填写 Remarks'],
    ['Error', '设备错误（含描述）', '可点击 Restart Service 或 Back 重试；必须填写 Remarks'],
    ['Not Done', '未执行扫描', '等待客户到场或转 Manual'],
], widths=[2.5, 4.0, 9.0])

P('说明：原始 FRS 中第四态命名为 Not Done，语义等同于「待执行」。客户反馈中提及的 Pending 与此为同一状态，本文档沿用原文命名 Not Done 以保持字段值一致。')

H(2, '9.4 设备支持与错误码映射')
T(['设备', '型号', '状态', '说明'], [
    ['Sagem MorphoSmart', '1350', '已纳入', '主力指纹扫描设备'],
    ['Dermalog', 'ZF1+', '已移除', '2015年11月经 PMO 确认移除'],
    ['SmartTag', 'FID688, CSmarT680', '已纳入', '便携式指纹扫描设备'],
], widths=[4.0, 4.0, 2.5, 5.0])

NEW('设备错误码映射表为新增，供开发对接设备 SDK 时使用。')
T(['设备', '设备状态码', '设备描述', '系统状态码', '系统描述'], [
    ['SmarTec', '0', 'Success', '0', 'Matched'],
    ['SmarTec', '-2', 'Read MyKad Failed', '-1', 'Error – Failed to read MyKad'],
    ['SmarTec', '-3', 'Fingerprint Verification Failed', '-2', 'Error – Fingerprint verification failed'],
    ['SmarTec', '-4 / -11 / 248 / 20001 / 20004 / 29000 / 其他', 'COM Port Error / Timeout / Card Not Inserted / Card Not Removed / User Quit / Create Object Error / Unknown', '-3', 'Error – Time Out'],
    ['SmarTec', '20002', 'IC Number Not Tally', '-4', 'Unmatched'],
    ['SmarTec', '20003', 'SmartTec Serial No Not Valid', '-5', 'Error – Device Serial Number Not Valid'],
    ['Sagem', '0', 'Match Successful', '0', 'Matched'],
    ['Sagem', '1', 'MyKad Not Found', '-1', 'Error – Failed to read MyKad'],
    ['Sagem', '2 / 3 / 11 / 12 / 15', 'Dongle / Reader / Scanner Not Found、Match Timeout', '-2', 'Error – Time Out'],
    ['Sagem', '14', 'Match Invalid Template', '-3', 'Error – Invalid template'],
    ['Sagem', '13 / 20002', 'Unmatched Thumbprint / IC Number Not Tally', '-4', 'Unmatched'],
], widths=[2.0, 3.5, 4.5, 2.0, 3.5])

H(2, '9.5 OTP 验证（RIB 侧实现）')
FIX('2026-09-01 评审确认：OTP 的触发、下发、输入、校验与重发全部由 RIB 负责，CrediOS 不实现相关界面与逻辑。本节参数仅作为对接口径记录。')

H(3, '9.5.1 OTP 参数')
T(['参数', '值', '说明'], [
    ['OTP 长度', '6位数字', '由 RIB 生成'],
    ['有效期', '10分钟', '从 OTP 发送时刻起计'],
    ['最大错误输入次数', '5次', '超过后转 Manual Acceptance'],
    ['成功记录', '终生有效', 'OTP 验证成功记录永久保存，不可删除'],
    ['触发时机', '仅第2次验证', '首次身份验证失败不可使用 OTP'],
], widths=[3.5, 3.0, 9.0])

TODO('P-18', 'OTP 重新发送的冷却时间待确认（建议 60 秒），以及单笔申请每日最大发送次数上限。该参数由 RIB 侧控制，需向对方确认。')

H(3, '9.5.2 CrediOS 与 RIB 的职责边界')
T(['职责', '归属'], [
    ['通知「第2次生物识别失败」', 'CrediOS → RIB（iqType = second bmt failed）'],
    ['展示 OTP 请求页、下发短信、收集并校验验证码、重发', 'RIB'],
    ['回传 OTP 验证结果', 'RIB → CrediOS（🔴 接口待确认，见 P-33）'],
    ['依据结果更新 Acceptance 状态', 'CrediOS'],
    ['等待期间的界面表现', 'CrediOS 显示等待态 + Refresh 按钮，不显示 OTP 输入框'],
], widths=[7.5, 8.0])

H(2, '9.6 验证异常汇总')
T(['验证类型', '异常场景', '处理逻辑'], [
    ['第1次 Biometric', '匹配失败', '允许重试（默认3次，可配置）'],
    ['第1次 Biometric', '失败达上限', '转 Manual Acceptance'],
    ['第1次 Biometric', '设备持续错误', '转 Manual Acceptance'],
    ['第2次 Biometric', '匹配失败', '转 OTP 验证'],
    ['第2次 Biometric', '15分钟超时', '转 OTP 验证'],
    ['OTP', '输入错误', '允许重新输入（上限5次）'],
    ['OTP', '10分钟过期', '允许重新发送'],
    ['OTP', '多次失败', '转 Manual Acceptance'],
    ['OTP', '短信网关异常', '重试3次后转 Manual Acceptance'],
    ['担保人 Biometric', '失败达上限', '转 Manual Acceptance（Hybrid 流程）'],
], widths=[3.5, 4.0, 8.0])

doc.add_page_break()

# =====================================================================
# 9 界面改造清单  ★NEW
# =====================================================================
H(1, '10. 界面改造清单')
NEW('本章为新增，对应原始 FRS HP075 第5章。v1.0 仅泛泛提及「列表页展示规则」，未指明需改造的具体界面及字段。')

H(2, '10.1 列表与查询类界面（3个）')
P('以下三个界面需新增 E-Acceptance 状态字段，取值与展示规则见第3.3节：')
T(['界面', '新增字段', '备注'], [
    ['To-Do List', 'E-Acceptance', '该界面为全用户共享，所有用户均可见此字段'],
    ['Pool List', 'E-Acceptance', '同上'],
    ['Application Inquiry', 'E-Acceptance', '同上'],
], widths=[4.0, 4.0, 7.5])

H(2, '10.2 CRA 类界面（2个）')
P('CRA Checker 与 CRA Maintenance 两个界面需新增/改造四个字段：')
T(['字段', '控件类型', '取值', '说明'], [
    ['Customer Biometric (Non-E)', 'Checkbox', 'Y / Blank', '由既有字段 Biometric 改名而来。当案件为 E-Acceptance 时该勾选框必须禁用（不适用）'],
    ['Identity Verification (E)', 'Display', 'Matched / Unmatched / Error（含描述）/ Not Done / N/A', 'N/A 表示案件走 Manual 流程'],
    ['Agreement Acceptance Type (E)', 'Display', 'Biometric / OTP / N/A', '客户最终采用的确认方式'],
    ['Status (E)', 'Display', 'Manual / Completed / Pending / Blank', 'Pending 与 Blank 仅在 Application Inquiry 中的 CRA 界面显示'],
], widths=[4.2, 2.0, 5.0, 4.3])

FIX('「Customer Biometric (Non-E) 改名 + E-Acceptance 时禁用」是原始 FRS 明确的改造要求，v1.0 完全未覆盖。')

H(2, '10.3 E-Acceptance 按钮')
T(['项目', '规则'], [
    ['位置', '申请页脚（footer）'],
    ['权限控制', '通过访问控制配置可见性，授权用户组：① EHP team（销售）② CRA team（复核）'],
    ['显示规则1', '只要存在 E-Acceptance 记录，按钮即 enable，即使状态已变更为 Manual'],
    ['显示规则2', '若案件为端到端纯 Manual Acceptance 案件，按钮不显示'],
    ['点击行为', '弹出 E-Acceptance 弹窗（含 Acceptance Log 表格与 Print Log 按钮）'],
], widths=[3.0, 12.5])

H(2, '10.4 E-Acceptance 弹窗')
T(['区域', '字段/元素'], [
    ['头部信息', 'Name（含ID号）、App. Ref. No.、Bank Name、Lending Type、Loan Account No.、SMS Infobip Status、Send Date'],
    ['操作按钮', 'Close、Print Log、Trigger Amendment SMS（状态规则见7.3节）'],
    ['Log 表格', '见第11.1节'],
], widths=[3.0, 12.5])

P('Bank Name 取值规则：Lending Type = Conventional 时显示「Hong Leong Bank Berhad」；Lending Type = Islamic 时显示「Hong Leong Islamic Bank」。')

doc.add_page_break()

# =====================================================================
# 10 文件管理与归档
# =====================================================================
H(1, '11. 文件管理与归档')

H(2, '11.1 两个归档位置的分工')
NEW('2026-09-01 评审明确了 Document Library 与文件中心（Attachment）的职责边界。')

T(['位置', '存放内容', '覆盖策略'], [
    ['Document Library', '**全部历史版本**：未签署的模板版本 + 已签署版本', '只增不覆盖，通过 History 查看与下载各版本'],
    ['文件中心 / Attachment', '**仅最新版本**', '每次电子签完成后自动覆盖（软覆盖，机制同 Application Form 的 override）'],
], widths=[4.0, 6.5, 5.0])

P('新增字段要求：Document Library 中的签约文书需增加一个 **标识字段（Tag）**，区分该版本为 `Template（未签署）` 还是 `Signed Version（已签署）`，便于按状态检索历史版本。', bold=True)

P('Manual Acceptance 的各版本文书同样全部保留在 Document Library。')

TODO('P-34', '「软覆盖」的技术含义需确认：文件中心展示最新版本，被覆盖的旧版本是否仍可从底层系统检索？评审倾向于软覆盖（可检索），需与开发确认实现方式。')

H(2, '11.2 电子签自动归档')
P('客户完成生物识别或 OTP 验证后，系统自动将已签署文件盖章后归档到 Attachment 附件页，CRA 可直接在系统内查看，无需柜员手动上传。自动挂载的5类文件见第5.3节 Step 7。')

TODO('P-35', '🔴 需向客户索取**带电子签名的文书样本（sample）**。当前理解是签名栏以「时间戳 + 姓名」替代手写签名，但需确认是否要求符合规范的 PDF 数字签名 / 电子印章。若需要，还需确认行内是否具备签章系统。该项影响文书生成方案。')

H(2, '11.3 归档状态追踪')
T(['状态', '说明'], [
    ['Pending', '待归档（文件已生成，等待归档至 Attachment）'],
    ['Archiving', '归档处理中'],
    ['Success', '归档成功'],
    ['Failed', '归档失败（需重试或人工介入）'],
], widths=[3.0, 12.5])

H(2, '11.4 版本控制与防覆盖')
T(['原则', '说明'], [
    ['原始文件永久保留', '首次签署完成的协议文件不可被覆盖，即使发生 Amendment 修订，原始文件仍保留'],
    ['修订文件独立版本', '每次 Amendment 生成的协议文件作为新版本独立存储，不覆盖任何历史版本'],
    ['Print Log 不可编辑', '生成后的 Print Log PDF 不可编辑、不可覆盖，仅可追加 Regenerate 记录'],
    ['操作日志不可删除', '所有 Acceptance Log 和系统状态变更日志不可删除、不可修改'],
], widths=[4.0, 11.5])

T(['规则', '说明'], [
    ['版本号生成', '初始版本为 v1.0，每次 Amendment 递增：v1.0 → v2.0 → v3.0'],
    ['版本命名', '{申请编号}_{文档类型}_v{版本号}.pdf，如 HP20260001_Agreement_v1.0.pdf'],
    ['版本关联', '所有版本通过申请编号关联，支持按申请编号查询所有历史版本'],
    ['变更摘要', '每个新版本自动生成变更摘要，记录变更字段和旧值/新值'],
    ['双向查询', '从申请可查询所有版本文件；从文件可查询所属申请'],
    ['版本对比', 'Sales 可在 CrediOS 中选择任意两个版本进行对比，差异字段高亮显示'],
], widths=[3.0, 12.5])

P('查看入口说明：Attachment 中展示最新版本；历史版本在文件中心（Document Library）入口查看。')

H(2, '11.5 签约类型双向打标机制')
NEW('本节为新增，对应原始 FRS HP075 5.11 与 Training Deck 2.4。v1.0 仅覆盖了四场景中的一个（第6.4节的纠错场景）。')

P('上传文档时，系统根据所选文档类型自动为案件打「签约类型」标签，共四个场景：')
T(['场景', '上传的文档类型', '案件当前类型', '系统行为'], [
    ['1', 'CRA HP Agreement (Manual)', '任意', '案件打标为 Manual Acceptance 类型'],
    ['2', 'CRA HP Agreement (E-Acceptance)', '任意', '案件打标为 E-Acceptance 类型'],
    ['3', 'CRA HP Agreement (Manual)', 'E-Acceptance', '弹窗二次确认「是否以 CRA HP Agreement (Manual) 继续？」→ 确认后案件由 E-Acceptance 改为 Manual Acceptance'],
    ['4', 'CRA HP Agreement (E-Acceptance)', 'Manual（误标纠错）', '弹窗二次确认 → 确认后系统自动抓取最新版 HP Agreement 挂载（Browse 按钮禁用，无需实际选择文件）→ 案件由 Manual 改回 E-Acceptance'],
], widths=[1.2, 4.5, 3.0, 6.8])

P('文档类型代码维护在公共代码表 UPLOAD_DOC_CHKLST，由 CS Admin 维护。')

H(2, '11.6 放款补件自动上传（Defect Rectification）')
NEW('本节为新增，对应原始 FRS HP075 第6章与 Training Deck 2.3。v1.0 完全未覆盖此功能模块。')

P('场景说明：CRA 复核时若发现文件缺失，会将案件打回至「Defects Pending Rectification」状态，Sales 需在 Defect Upload Document 页面补齐文件。由于电子签流程中签署文件已自动归档，补件时不应要求 Sales 重复手动上传。')

T(['字段', '控件类型', '必填', '说明'], [
    ['Document Type', 'Dropdown', 'M', '选择需补充的文档类型'],
    ['Upload Manual', 'Radio', 'M', 'Yes = 手动上传附件；No = 系统自动挂载（默认）'],
    ['File', 'File Browser', 'O', 'Upload Manual = Yes 时必填并按既有逻辑上传；= No 时该字段不强制，Browse 按钮禁用'],
    ['Description', 'Input', 'O', '系统根据挂载文件预填描述，字段保持可编辑'],
    ['Submit', 'Button', '—', '提交后系统按 Upload Manual 取值执行自动挂载或手动上传'],
], widths=[3.0, 2.5, 1.5, 8.5])

P('自动挂载规则：', bold=True)
BULLET('系统抓取所选文档类型的最新版本挂载')
BULLET('若该案件发生过 Amendment，系统自动抓取修订后的最新版本')
BULLET('若系统内不存在该文档类型的文件，提示「document not exist, please select the document and upload」，此时 Sales 可将 Upload Manual 切换为 Yes 走手动上传流程')

H(2, '11.7 Manual Acceptance 文件强制上传')
P('详见第6.3节六类材料清单。')

doc.add_page_break()

# =====================================================================
# 11 签约日志与报告输出
# =====================================================================
H(1, '12. 签约日志与报告输出')

H(2, '12.1 Acceptance Log')

H(3, '12.1.1 三种记录类型')
T(['类型', '英文', '触发时机'], [
    ['初次签约', 'Hirer Initial Acceptance', '客户首次完成签约'],
    ['修订签约（第N次）', 'Hirer Amendment Acceptance (1st/2nd/3rd…)', '客户通过 e-Tracker 完成修订确认，按顺序编号'],
    ['银行签约', 'Bank Acceptance', '银行方完成签约确认并放款时盖章'],
], widths=[3.5, 6.0, 6.0])

P('支持多轮 Amendment，Times of Amendment 字段显示第几次。')

H(3, '12.1.2 Log 表格字段')
T(['字段', '取值', '说明'], [
    ['E-Acceptance Type', '见 11.1.1', '签约动作类型'],
    ['Product Disclosure Sheet', 'Y / Blank', 'Y = 客户已对该文件执行接受；Blank = 该文件不适用于本次签约类型'],
    ['Second Schedule Part 1', 'Y / Blank', '同上（客户或银行执行）'],
    ['HP Agreement & T&C', 'Y / Blank', '同上（客户或银行执行）'],
    ['Appendix 4', 'Y / Blank', '同上，仅二手/翻新车适用'],
    ['E-Acceptance Status', 'Completed', '该轮签约已在客户端完成'],
    ['E-Acceptance Date Time', '格式：20/12/2021 11:45 AM', 'Bank Acceptance 行显示放款执行时间'],
    ['Remark', '见下', '按类型显示：Initial 显示 Biometric Acceptance 或 OTP Acceptance；Amendment 不显示；Bank Acceptance 显示 Funded'],
], widths=[4.5, 3.5, 7.5])

H(3, '12.1.3 Log 查询功能')
T(['字段', '说明'], [
    ['入口', 'E-Acceptance 按钮（权限规则见第9.3节）'],
    ['列表字段', '申请编号、客户姓名、ID号、Acceptance Mode、Acceptance Status、Biometric状态、OTP状态、Sales、操作时间'],
    ['导出', '支持导出'],
    ['不可变性', 'Acceptance Log 记录不可修改、不可删除，每次重新签约产生新记录'],
], widths=[3.0, 12.5])

H(2, '12.2 系统状态变更日志')
T(['事件类型', '记录内容'], [
    ['Acceptance Mode 变更', '变更前值、变更后值、Sales、时间戳'],
    ['Acceptance Status 变更', '变更前值、变更后值、Sales、时间戳'],
    ['Biometric 结果记录', '验证类型、结果、尝试次数、时间戳'],
    ['OTP 发送/验证记录', '发送时间、验证结果、有效期'],
    ['文件归档事件', '文档名称、归档状态、时间戳'],
    ['文件版本变更', '旧版本号、新版本号、变更摘要'],
], widths=[4.5, 11.0])

H(2, '12.3 E-Acceptance Print Log PDF')
P('从 E-Acceptance 弹窗点击「Print Log」按钮生成，按申请（application）为单位输出，是 Acceptance Log 表格的可打印/归档静态版本。')

P('术语说明：客户反馈中提及的「E-Acceptance Report」即本报表，原始 FRS 正式名称为 E-Acceptance Print Log。当案件处于 Pending E-Acceptance 状态时，Print Log 内容即为 E-Acceptance Report。')

H(3, '12.3.1 Header 字段')
T(['字段', '说明'], [
    ['Name', '申请人姓名（含 ID 号）'],
    ['Relationship to Application', '申请人与本申请的关系'],
    ['App. Ref. No.', '申请编号'],
    ['Bank Name', '按 Lending Type 判断：Conventional → Hong Leong Bank Berhad；Islamic → Hong Leong Islamic Bank'],
    ['Lending Type', '贷款类型'],
    ['Loan Account No.', '贷款账号'],
    ['Generated Date', '生成日期'],
    ['Generated By', '生成人（Sales ID）'],
], widths=[4.5, 11.0])

FIX('v1.0 的 Header 字段缺 Relationship to Application、Bank Name、Lending Type、Loan Account No. 四项（客户反馈第10条明确要求），本版恢复并保留 v1.0 新增的 Generated Date / Generated By。')

H(3, '12.3.2 明细字段')
T(['字段', '说明'], [
    ['E-Acceptance Type', '签约动作类型（Initial / Amendment(N) / Bank Acceptance）'],
    ['Product Disclosure Sheet', 'Y / Blank'],
    ['Second Schedule Part 1', 'Y / Blank'],
    ['HP Agreement & T&C', 'Y / Blank'],
    ['Appendix 4', 'Y / Blank'],
    ['E-Acceptance Status', 'Completed'],
    ['E-Acceptance Date Time', '执行时间'],
    ['Remark', '按类型显示不同内容，见 11.1.2'],
    ['Signing Method', 'E-signing / Paper Signing'],
    ['Biometric Result', '生物识别结果'],
    ['OTP Result', 'OTP 结果（如适用）'],
], widths=[4.5, 11.0])

FIX('v1.0 将明细改为「按文档看签署详情」，丢失了原文四列 Y/Blank 文档接受标记。这四列是审计时核对「本轮签约客户接受了哪几份文件」的依据，也是区分 Hirer 行与 Bank Acceptance 行的方式（银行只签 Second Schedule 与 HP Agreement，不签 PDS）。本版恢复四列并与 v1.0 新增字段并存。')

TODO('P-28', 'Print Log 明细字段最终采用哪一版需确认：本版为「原文四列 Y/Blank + v1.0 新增字段」的合并方案，字段数较多。若客户偏好精简，需确定取舍。')

H(3, '12.3.3 生成与重生成规则')
T(['规则', '说明'], [
    ['自动生成', 'Acceptance Status 变更为 Completed 时，系统自动生成 Print Log PDF'],
    ['手动 Regenerate', 'Sales 可在 Print Log 界面点击「Regenerate」重新生成 PDF'],
    ['手动 Reprint', 'Sales 可点击「Reprint」重新打印 PDF'],
    ['防篡改', '生成后的 PDF 不可编辑，确保审计完整性；仅可追加 Regenerate 记录'],
], widths=[3.5, 12.0])

H(2, '12.4 Biometric Report')
T(['字段', '说明'], [
    ['入口', 'Biometric 详情页 / 应用内 Biometric Tab 的「Report」按钮'],
    ['生成粒度', '按申请人生成。一个案件若有多个验证对象（如主申请人 + 担保人），生成多份独立报表'],
    ['单份内容', '同一申请人的多种验证类型（Identity Verification / E-Acceptance Confirmation）合并在一份报表内，每种类型一行'],
    ['重试处理', '多次重试不逐条列出，时间字段取最新一次，累计次数体现在 Triggering Counter'],
    ['字段', '见第4.3节 Step 3 的详情页字段清单'],
    ['归档', '生成后自动归档至 Attachment，可在文件中心下载 PDF；多次生成的历史版本在 Document Library 查看，Attachment 中仅保留最新版本'],
], widths=[3.0, 12.5])

P('说明：CrediOS 内完成的生物识别（无论 E-Acceptance 或 Manual Acceptance），其报表由系统自动归档，不需要柜员手动上传，CRA 可直接在系统内查看。')

H(2, '12.5 CRA Report')
P('在申请界面点击「Report」按钮下载。新增列：Identity Verification (E)、Agreement Acceptance Type (E)、Status (E)。既有字段 Customer Biometric (Non-E) 在 E-Acceptance 案件下显示为 Blank。')

H(2, '12.6 HP Funding Report')
P('新增三个字段：Identity Verification (E)、E-Acceptance Status (E)、Agreement Acceptance Type (E)。Customer Biometric (Non-E) 在 E-Acceptance 案件下显示为 Blank。')

doc.add_page_break()

# =====================================================================
# 12 工作流校验与错误提示  ★NEW
# =====================================================================
H(1, '13. 工作流校验与错误提示')
NEW('本章为新增，对应原始 FRS HP075 12.3 与 HP047 第6章。沿用原文英文错误文案，可显著降低一线人员的适应成本。')

H(2, '13.1 校验器清单')
T(['校验器', '触发条件', '错误提示文案'], [
    ['ValidateEAcceptanceEligible', '不符合电子签准入条件却尝试走电子签', 'Application not eligible for E-Acceptance, please proceed with manual acceptance'],
    ['ValidateEAcceptanceCombo', 'E_ACCEPTANCE_COMBO=E[E] 时尝试改为人工签', 'Application not allow to change to manual acceptance, please proceed with E-Acceptance'],
    ['ValidateBiometricComplete', '生物识别状态为 Unmatched', 'Verification failed, please proceed with manual acceptance'],
    ['ValidateBiometricComplete', '生物识别状态为 Not Done', 'Please complete the [Verification type] Biometric Verification（[] 内代入 Identity Verification 或 E-Acceptance Confirmation）'],
    ['ValidateEAcceptanceComplete', '未收到 RIB 回传的签约确认标志', 'Customer has not performed acceptance.'],
    ['isEAccptAmendmentComplete', '未收到客户的修订确认回执', 'Customer has not performed acceptance on the amendment request.'],
    ['MANDATORY CHECK', '命中强制矩阵但生物识别未完成', 'Please complete the Biometric scanning'],
    ['UNSUCCESSFUL BIOMETRIC', 'Status 为 Unmatched 或 Error 但 Remarks 为空', 'Biometric Status is Unmatched/Error without any remarks. Please input the remarks in the Biometric Details Screen.'],
    ['MAX RETRY', '超过 BIOMETRIC_ATTEMPT_COUNTER 配置值', 'Exceed maximum retry for biometric scanning'],
    ['SUCCESSFUL COUNTER', '超过 BIOMETRIC_SUCCESSFUL_COUNTER 配置值', 'Exceed the number of successful attempt for biometric scanning'],
], widths=[4.0, 4.5, 7.0])

H(2, '13.2 两个计数器的区别')
T(['计数器', '统计范围', '用途'], [
    ['BIOMETRIC_ATTEMPT_COUNTER', '所有尝试，含 Error', '限制总重试次数，防止无限重试'],
    ['BIOMETRIC_SUCCESSFUL_COUNTER', '仅 Matched / Unmatched，不含 Error', '限制有效验证次数，防止对已完成的验证重复触发'],
], widths=[5.0, 5.0, 5.5])

FIX('v1.0 将 Successful Counter 描述为「用于统计」，弱化了其拦截作用。该计数器超限时同样会阻断流程。')

H(2, '13.3 校验器配置位置')
T(['校验器', '配置的工作流节点'], [
    ['MANDATORY CHECK', '审批决策之后、主机接口之前'],
    ['UNSUCCESSFUL BIOMETRIC（Remarks 校验）', '签约步骤（Letter Printing 对应节点）'],
    ['ValidateEAcceptanceEligible / Combo', '签约模式选择时'],
    ['ValidateEAcceptanceComplete', '案件流出签约步骤时'],
    ['isEAccptAmendmentComplete', '案件流出签约步骤时（仅针对提交过 CILT 的案件）'],
], widths=[6.0, 9.5])

doc.add_page_break()

# =====================================================================
# 13 配置参数汇总  ★NEW
# =====================================================================
H(1, '14. 配置参数汇总')
NEW('本章为新增。v1.0 仅列出3个生物识别参数，缺4个流程级开关。')

H(2, '14.1 流程级开关')
T(['Constant Code', '作用', '取值', '缺失影响'], [
    ['E_ACCEPTANCE_ON_OFF', '电子签总开关', 'ON / OFF', 'OFF 时全部案件走人工签路径；用于灰度上线与紧急回退'],
    ['E_ACCEPTANCE_GOLIVE_DATE', '生效日期', 'DD/MM/YYYY 00:00', '仅对该日期之后创建的新案件生效，存量/在途案件继续使用旧流程与旧界面'],
    ['E_ACCEPTANCE_COMBO', '签约方式切换策略', 'E[E] / E[M]', 'E[E]=只能电子签，不允许改人工签；E[M]=允许电子签、人工签或两者组合'],
    ['BIOMETRIC_CHECK_SELLER', '卖家生物识别开关', 'I / C / OFF', 'I=个人卖家需验证；C=企业卖家需验证；OFF=关闭。开启后即为强制完成项'],
], widths=[4.5, 3.0, 3.0, 5.0])

P('重要提示：本文档中大量描述的「失败后转 Manual Acceptance」降级路径，其可行性取决于 E_ACCEPTANCE_COMBO 的取值。若配置为 E[E]，系统将拦截转人工签的操作并提示错误（见第12.1节）。')

TODO('P-29', 'E_ACCEPTANCE_COMBO 的生产环境取值待确认。若采用 E[E]，则本文档所有「转 Manual」的异常处理路径均需重新设计。')

TODO('P-30', '存量/在途案件的处理方式待确认：上线时已进入签约流程但未完成的案件，是继续走旧流程，还是需要数据迁移到新流程？')

H(2, '14.2 生物识别参数')
T(['Constant Code', '作用', '建议值'], [
    ['BIOMETRIC_ATTEMPT_COUNTER', '生物识别最大尝试次数（含 Error）', '3（系统上限999）'],
    ['BIOMETRIC_SUCCESSFUL_COUNTER', '最大有效验证次数（不含 Error）', '3（系统上限999）'],
    ['BIOMETRIC_MANDATORY_CHECK', '生物识别是否强制（矩阵，见2.4.3）', '可配置'],
    ['BIOMETRIC_CRITERIA_CHECK', '生物识别准入矩阵（见2.4.2）', '可配置'],
], widths=[5.0, 6.5, 4.0])

H(2, '14.3 e-Tracker 参数')
T(['参数', '说明', '建议值'], [
    ['前签约天数', '审批通过后允许签约的最大天数', '30天'],
    ['后签约天数', '签约后允许放款的最大天数', '90天'],
    ['取消/拒绝天数', '取消/拒绝后记录保留天数', '180天'],
    ['放款年限', '放款后记录可查询年限', '7年'],
], widths=[3.5, 8.0, 4.0])

TODO('P-14', '上述 e-Tracker 参数的生产环境实际配置值待客户确认。')

doc.add_page_break()

# =====================================================================
# 14 Letter Template  ★NEW
# =====================================================================
H(1, '15. Letter Template 与文档生成规则')
NEW('本章为新增。v1.0 将此项列为待确认 P-22，实际原始 FRS HP075 13.2.1 与 13.4 已完整定义，现予补齐并结案。')

H(2, '15.1 模板清单（8个）')
P('维护位置：Parameter > Common Codes > Code Table: DOCUMENT MANAGEMENT SYSTEM，代码类型 LETTER_T。')
T(['#', '代码', '模板名称', '用途'], [
    ['1', 'HPPDS_E', 'PRODUCT DISCLOSURE SHEET (E-ACCEPTANCE)', '展示电子签名信息'],
    ['2', 'HPSS_E', 'HP SECOND SCHEDULE (E-ACCEPTANCE)', '展示电子签名信息'],
    ['3', 'HPA_E', 'HP AGREEMENT (PART 1) (E-ACCEPTANCE)', '展示电子签名信息'],
    ['4', 'HPA4_E', 'APPENDIX 4 CUSTOMER/HIRER INDEMNITY LETTER (E-ACCEPTANCE)', '展示电子签名信息'],
    ['5', 'HPPDS_EA', 'PRODUCT DISCLOSURE SHEET (E-ACCEPTANCE AMENDMENT)', '展示修订高亮'],
    ['6', 'HPSS_EA', 'HP SECOND SCHEDULE (E-ACCEPTANCE AMENDMENT)', '展示修订高亮'],
    ['7', 'HPA_EA', 'HP AGREEMENT (PART 1) (E-ACCEPTANCE AMENDMENT)', '展示修订高亮'],
    ['8', 'HPA4_EA', 'APPENDIX 4 CUSTOMER/HIRER INDEMNITY LETTER (E-ACCEPTANCE AMENDMENT)', '展示修订高亮'],
], widths=[1.0, 2.2, 8.3, 4.0])

H(2, '15.2 签名字段替代规则')
P('电子签场景下，原纸质协议的手写签名栏改为显示时间戳，这是电子签呈现的核心逻辑：')
T(['模板', '字段', '替代内容', '示例'], [
    ['HP Agreement Part 1', 'Signature of Hirer(s)', '客户执行签约的日期时间', '25/11/2021 11:25:20 AM'],
    ['HP Agreement Part 1', 'Name', '主申请人姓名', 'Tan Ming Lee'],
    ['HP Agreement Part 1', 'NRIC No.', '主申请人身份证号', '970913-10-5621'],
    ['HP Agreement Part 1', 'Date', '客户执行签约的日期', '25/11/2021'],
    ['HP Agreement Part 1', 'Bank Authorised Signature', '银行接受该笔业务的日期时间', '25/11/2021 11:25:20 AM'],
    ['HP Agreement Part 1', 'Funding Date', '放款日期', 'the 24th day of November 2021'],
    ['Product Disclosure Sheet', "Applicant's Signature", '客户执行签约的日期时间', '25/11/2021 11:25:20 AM'],
    ['HP Second Schedule', 'Signature of Dealer / Owner', '申请创建日期时间（推送前端前先盖章）', '25/11/2021 11:25:20 AM'],
    ['HP Second Schedule', 'Signature of Hirer', '客户执行签约的日期时间', '25/11/2021 11:25:20 AM'],
    ['Appendix 4', 'Signature of Hirer(s)', '客户执行签约的日期时间', '25/11/2021 11:25:20 AM'],
    ['Appendix 4', 'Date（Header）', '申请创建日期时间', '25/11/2021 11:25:20 AM'],
], widths=[3.8, 4.0, 5.0, 2.7])

H(2, '15.3 排版合规规则')
P('以下两条规则涉及《1967年租购法令》（Hire Purchase Act 1967）合规，不可遗漏：')
T(['规则', '说明'], [
    ['移除见证人签名栏', 'E-Acceptance 版本的 HP Agreement Part 1 必须移除 Witness signature 栏位'],
    ['条款加删除线', 'E-Acceptance 版本的 HP Second Schedule PART I 中，Subparagraph 4(1)(b)(i) 必须加删除线（strikethrough）'],
], widths=[4.5, 11.0])

H(2, '15.4 模板语言版本')
T(['模板', '语言版本'], [
    ['HP Agreement Part 1', 'Hire Purchase Agreement Part 1（英/马）、Hire Purchase i-Agreement Part 1（英/马），共4个版本'],
    ['Product Disclosure Sheet', 'HP/IHP PDS Conventional（英/马）、HP PDS Islamic（英/马），共4个版本'],
    ['HP Second Schedule', 'Hire Purchase Second Schedule PART 1'],
    ['Appendix 4', 'Appendix 4 Customer or Hirer Indemnity Letter'],
], widths=[4.0, 11.5])

doc.add_page_break()

# =====================================================================
# 15 定时任务  ★NEW
# =====================================================================
H(1, '16. 定时任务与批处理')
NEW('本章为新增，对应原始 FRS HP075 第11章。v1.0 完全未覆盖，其中 LHDN 印花税申报为监管合规产物。')

H(2, '16.1 定时任务清单')
T(['任务名', '用途', '改造内容', '执行频率'], [
    ['executeAutoCRAEmailTrigger', 'CRA 邮件触发', 'e-HP Agreement 合并逻辑迁移至 generateCRACompile 工作流函数', '每5分钟，及 01:00–05:00 整点'],
    ['executeGenerateLHDNXMLFileJobTrigger', '生成 LHDN 印花税申报 XML', '按签约类型拆分生成，详见15.2节', '每日 06:30'],
    ['batchSendSmsTrigger', '批量短信发送', '排除走 Infobip 网关的 Amendment 通知短信，避免重复发送', '每30分钟，08:00–19:59'],
], widths=[4.5, 3.0, 5.5, 2.5])

H(2, '16.2 LHDN 印花税申报文件拆分')
P('合规背景：电子签约改变了协议的存在形态（数字签名时间戳替代手写签名），向马来西亚税务局（LHDN）提交的印花税申报文件组装方式必须相应调整。')

T(['签约类型', 'XML 内容', '说明'], [
    ['E-Acceptance', '数字协议 + E-Acceptance Report', '包含电子签署的 HP Agreement & T&C、Appendix、LHDN log；附件以 base64 编码嵌入'],
    ['Manual Acceptance', '按既有 BAU 逻辑生成', '沿用现有申报文件结构'],
], widths=[3.5, 5.0, 7.0])

P('批量提交要求：电子签与非电子签案件的批量提交文本文件需分开生成，不可混在同一批次。')

TODO('P-31', 'LHDN XML 在 CrediOS 新架构下的生成方式待确认：是沿用定时任务批量生成，还是改为签约完成后实时生成？申报报文的字段结构是否有变化？')

doc.add_page_break()

# =====================================================================
# 16 接口
# =====================================================================
H(1, '17. 系统集成与接口')

H(2, '17.1 e-Tracker / RIB 查询类接口')
T(['#', '接口名称', '事务码', 'EAI事务码', '方向'], [
    ['1', 'Validate Customer Inquiry', 'VALID_CUST_INQ', '—', 'eTracker/RIB → CrediOS'],
    ['2', 'Application Status Inquiry', 'APL_STAT_INQ', '423001', 'eTracker/RIB → CrediOS'],
    ['3', 'Application Detail Inquiry', 'APL_DTL_INQ', '—', 'eTracker/RIB → CrediOS'],
    ['4', 'Attachment Inquiry', 'ATTACH_INQ', '—', 'eTracker/RIB → CrediOS'],
    ['5', 'Document Inquiry', 'EHPA_ATTACH_INQ', 'New', 'CrediOS → RIB'],
    ['6', 'Download Attachment Inquiry', 'DLOAD_ATTACH_INQ', 'New', 'CrediOS → eTracker/RIB'],
], widths=[1.0, 5.0, 3.5, 2.0, 4.0])

H(2, '17.2 签约结果回传接口')
T(['#', '接口名称', '事务码', 'EAI事务码', '方向'], [
    ['7', 'Acceptance Decision Message', 'LD0008', '421015', 'RIB/eTracker → CrediOS'],
    ['8', 'RIB Service（生物识别与签约状态更新）', '—', 'New', 'CrediOS → RIB'],
], widths=[1.0, 5.0, 3.5, 2.0, 4.0])

H(2, '17.3 hpAgreementInd 复合状态字段')
NEW('本节为新增，对应原始 FRS HP075 10.2 与 12.1.1。该字段决定 e-Tracker 前端展示什么进度，v1.0 未覆盖。')

P('Application Status Inquiry（APL_STAT_INQ）报文新增 hpAgreementInd 字段，由四个内部指示器组合推算：')
T(['eAcceptance Indicator', '1st Signing Complete', 'Ready for 2nd Signing', '2nd Signing Complete', 'eTracker Status', 'hpAgreementInd'], [
    ['Y', 'Y', '空', '空', 'APPR1 或 APPR2', '1'],
    ['Y', 'Y', 'Y', '空', 'APPR1 或 APPR2', '2'],
    ['Y', 'Y', 'Y', 'Y', 'APPR1 或 APPR2', '3'],
    ['空', 'Y', '空', '空', 'APPR1 或 APPR2', '4'],
    ['空', 'Y', 'Y', '空', 'APPR1 或 APPR2', '5'],
    ['空', 'Y', 'Y', 'Y', 'APPR1 或 APPR2', '6'],
], widths=[2.8, 2.6, 2.8, 2.6, 2.5, 2.2])

H(3, '17.3.1 四个指示器的盖章逻辑')
T(['指示器', '电子签案件', '非电子签案件'], [
    ['eAcceptance Indicator', '有电子签标记时置 Y', '置空'],
    ['1st Signing Complete', '首次触发签约日期盖章时置 Y。注意：需在客户完成第2次生物识别时才置 Y；若第2次生物识别失败走了 OTP，则在收到 RIB 的 OTP 成功消息时置 Y', '首次触发签约日期盖章时置 Y'],
    ['Ready for 2nd Signing', '用户在签约步骤触发「Ready for Amendment Acceptance」动作时置 Y', '案件流出协议文件生成环节且签约日期非空时置 Y'],
    ['2nd Signing Complete', 'AEM/RIB 返回修订/二次签署成功（Decision=AC + 2nd Signing Flag），且 Ready for 2nd Signing=Y 时置 Y', '案件流出签约步骤且签约日期非空、Ready for 2nd Signing=Y 时置 Y'],
], widths=[3.5, 6.0, 6.0])

P('实现提示：1st Signing Complete 不是在第一次生物识别后盖章，而是在整个签约动作闭环（第2次生物识别或 OTP 成功）后才盖章，与字面直觉相反，需特别注意。')

TODO('P-32', '新的 Acceptance Mode + Status 状态模型与上述四个指示器的映射关系待确认。e-Tracker 展示依赖这四个指示器，需明确新模型如何推导出这四个值。')

H(2, '17.4 e-Tracker AEM 前端 API')
T(['#', 'API 端点', '用途'], [
    ['9', 'etrackapplicationstatuslisting', '列表页 — 获取申请状态列表'],
    ['10', 'etrackapplicationstatusdetails', '详情页 — 获取 Facility Details'],
], widths=[1.5, 6.5, 7.5])

TODO('P-15', 'AEM e-Tracker Field Mapping 完整字段映射待补充。')

H(2, '17.5 短信网关接口')
T(['字段', '说明'], [
    ['网关提供商', 'Infobip'],
    ['触发时机', '① Amendment 修订通知；② OTP 验证码下发'],
    ['路由规则', 'Gateway=Infobip 的模板走 Infobip；Gateway=Exchange 或空值的模板走 M3Tech 网关'],
    ['日志表', 'SYS_INFOBIP_LOG'],
], widths=[3.0, 12.5])

TODO('P-09', 'RIB ↔ CrediOS 之间的接口协议细节（消息格式、安全机制、字段级定义）待补充，原始 FRS 中该部分内容在 TRS 文档，需另行索取。')

doc.add_page_break()

# =====================================================================
# 17 客户通知与状态查询
# =====================================================================
H(1, '18. 客户通知与状态查询')

H(2, '18.1 短信通知')
T(['模板', '说明', '变量'], [
    ['首次签约通知', '通知客户进入 RIB 完成 E-signing 或前往网点 Paper Signing', '申请编号、eTracker 链接'],
    ['修订签约通知', '通知客户协议已修订，需重新签署（模板配置见7.6节）', '申请编号、eTracker 链接'],
    ['OTP 验证码', '包含6位数字 OTP', 'OTP 验证码、有效期（10分钟）'],
], widths=[3.5, 7.5, 4.5])

TODO('P-10', '电子签约涉及的所有短信模板完整清单待补充（含马来语版本、发送时机、字符数限制）。')

H(2, '18.2 手机号获取')
P('系统通过 Validate Customer Inquiry 接口（VALID_CUST_INQ）按客户 ID 查询手机号列表。若客户有多个手机号，Sales 需确认选择正确号码。')

H(2, '18.3 e-Tracker 客户端')

H(3, '18.3.1 功能概述')
P('e-Tracker 是客户查询签约进度的在线平台，基于 AEM（Adobe Experience Manager）构建，客户通过 www.hlb.com.my/etracker 访问。除查询进度外，还承载 Amendment 远程重签功能（见7.3节）。')

H(3, '18.3.2 签约阶段展示')
T(['阶段', '说明', '状态指示'], [
    ['Application Submission', '申请提交', '灰/黄/绿三态'],
    ['Application Decision', '贷款决策', '灰/黄/绿三态'],
    ['Agreement Signing', '协议签署', '未开始灰色；进行中黄色；已完成绿色'],
    ['Disbursement', '放款', '灰/黄/绿三态'],
], widths=[4.0, 4.0, 7.5])

P('客户在 Agreement Signing 阶段变绿后，可下载 Product Disclosure Sheet 副本。')

H(3, '18.3.3 e-Tracker 状态码')
T(['状态码', '含义', '说明'], [
    ['PEND', '申请提交中', '申请正在处理'],
    ['PEND2', '上诉中', '待上诉决定'],
    ['APPR1', '贷款决策（无变化）', '审批通过，条件未变'],
    ['APPR2', '贷款决策（有变化）', '审批通过，条件有变'],
    ['FUND', '已放款', '贷款已放款'],
    ['CANC', '已取消', '申请已取消'],
    ['REJ', '已拒绝', '申请被拒绝'],
    ['NA', '未找到', '状态不明确或超期'],
], widths=[2.5, 5.0, 8.0])

H(3, '18.3.4 「Not Found」触发条件')
T(['触发条件', '说明'], [
    ['超时未签约', '超过配置的前签约天数（见13.3节）'],
    ['超时未放款', '超过配置的后签约天数'],
    ['取消/拒绝后超期', '超过配置的取消/拒绝天数'],
    ['放款超期', '超过配置的放款年限'],
    ['草稿状态', '申请处于草稿阶段'],
], widths=[4.0, 11.5])

doc.add_page_break()

# =====================================================================
# 18 异常分支汇总
# =====================================================================
H(1, '19. 异常分支汇总')

H(2, '19.1 主流程异常（A–F）')
T(['编号', '异常场景', '影响流程', '处理逻辑', '状态变更'], [
    ['A-1', '第1次 Biometric 失败（Unmatched）', 'Full/Hybrid', '允许重试（默认3次，可配置）', 'Mode 不变, Status=Pending'],
    ['A-2', '第1次 Biometric 失败达上限', 'Full/Hybrid', '转 Manual Acceptance', 'Mode=Manual Acceptance, Status=Pending'],
    ['A-3', 'Biometric 设备错误', 'Full/Hybrid', '可点 Restart Service 或 Back 重试；持续错误转 Manual', '保持 / 转 Manual'],
    ['A-4', '客户未到场（Not Done）', 'Full/Hybrid', 'Sales 可选择等待或转 Manual', 'Status=Pending'],
    ['B-1', '客户未在 RIB 完成 E-signing', 'Full', '系统等待，不自动超时；Sales 可电话提醒', 'Mode=E-Acceptance, Status=Pending'],
    ['B-2', 'RIB 系统不可用', 'Full', '等待恢复；持续不可用则转 Manual', '保持 / 转 Manual'],
    ['B-3', 'RIB 回传 Acceptance Flag 超时', 'Full', '系统告警，Sales 介入判断', 'Status=Pending'],
    ['C-1', '第2次 Biometric 失败或超时', 'Full', '转 OTP 验证', 'Status=Pending'],
    ['C-2', 'OTP 发送失败（网关异常）', 'Full', '重试3次；持续失败转 Manual', '保持 / 转 Manual'],
    ['C-3', 'OTP 输入错误', 'Full', '允许重新输入（上限5次）', 'Status=Pending'],
    ['C-4', 'OTP 过期（10分钟）', 'Full', '允许重新发送', 'Status=Pending'],
    ['C-5', 'OTP 多次失败', 'Full', '转 Manual Acceptance', 'Mode=Manual Acceptance, Status=Pending'],
    ['C-6', '客户手机号不可用', 'Full', 'Sales 确认手机号，必要时更新后重新发送', 'Status=Pending'],
    ['D-1', '客户未注册 RIB', 'Full', '第1次 Biometric 成功后发现 → 转 Manual', 'Mode=Manual Acceptance, Status=Pending'],
    ['D-2', '客户 RIB 账户被冻结/停用', 'Full', '无法登录 RIB → 转 Manual', 'Mode=Manual Acceptance, Status=Pending'],
    ['E-1', 'eTracker 返回 Not Found', '全部', 'Sales 介入排查', '保持原状态'],
    ['E-2', 'eTracker 状态码为 CANC/REJ', '全部', '签约流程终止', '终止'],
    ['E-3', 'eTracker 状态与 CrediOS 不一致', '全部', '需人工排查', '保持原状态'],
    ['F-1', 'CrediOS 系统宕机', '全部', '恢复后根据 Mode + Status 字段恢复流程', '保持原状态'],
    ['F-2', 'Sales 误操作', '全部', '可手动修正状态（需权限控制）', '手动修正'],
    ['F-3', '网络中断', '全部', '设备脱机或 RIB 连接中断，恢复后重试', '保持原状态'],
], widths=[1.2, 4.0, 2.0, 4.5, 3.8])

TODO('P-16', '客户未注册 RIB（D-1）时的引导流程待确认：是否需要在签约现场引导客户开通网银？开通后能否当场转回 E-Acceptance？')

TODO('P-05', '设备异常与网络中断的详细处理流程待补充：包括断点续传机制、超时阈值、告警通知对象。')

H(2, '19.2 Hybrid 特有异常')
T(['编号', '异常场景', '处理逻辑', '状态变更'], [
    ['H-1', '担保人 Biometric 失败达上限', '转 Manual Acceptance', 'Mode=Manual Acceptance, Status=Pending'],
    ['H-2', '担保人不在场', 'Sales 可选择等待或转 Manual', 'Mode=Manual Acceptance, Status=Pending'],
    ['H-3', '客户拒绝签署纸质协议', 'Sales 记录原因，申请进入异常处理', 'Mode=Manual Acceptance, Status=Pending'],
    ['H-4', '纸质签字部分签署', '仅主申请人签署、担保人未签署 → 不视为完成', 'Mode=Manual Acceptance, Status=Pending'],
    ['H-5', '纸质文件遗失', '重新打印协议文件', 'Mode=Manual Acceptance, Status=Pending'],
    ['H-6', '扫描文件上传失败', '允许重新上传', 'Mode=Manual Acceptance, Status=Pending'],
], widths=[1.2, 4.5, 5.0, 4.8])

FIX('本表 Mode 取值 v1.0 中误写为 E-Acceptance，本版修正为 Manual Acceptance。')

H(2, '19.3 Manual 特有异常')
T(['编号', '异常场景', '处理逻辑', '状态变更'], [
    ['M-1', '文件上传不完整', '提示补充，不允许完成签约', 'Mode=Manual Acceptance, Status=Pending'],
    ['M-2', '文件格式错误', '提示格式要求，允许重新上传', 'Mode=Manual Acceptance, Status=Pending'],
    ['M-3', '文件大小超限', '提示限制，允许压缩后重新上传', 'Mode=Manual Acceptance, Status=Pending'],
    ['M-4', '外籍客户 Biometric 豁免', '系统自动标记豁免', 'Mode=Manual Acceptance, Status=Pending'],
    ['M-5', 'Non-Individual 无 Biometric', '系统跳过 Biometric 步骤', 'Mode=Manual Acceptance, Status=Pending'],
], widths=[1.2, 4.5, 5.0, 4.8])

TODO('P-12', '附件大小限制的具体值待确认（单文件上限、单案件累计上限、支持的文件格式）。')

H(2, '19.4 辅助功能异常')
T(['编号', '异常场景', '处理逻辑'], [
    ['AUX-1', 'EDMS 归档失败', '归档状态设为 Failed，允许重试或人工介入'],
    ['AUX-2', 'Print Log PDF 生成失败', '允许手动 Regenerate'],
    ['AUX-3', '版本对比服务不可用', '提示稍后重试'],
    ['AUX-4', 'Acceptance Log 导出失败', '允许重新导出'],
    ['AUX-5', 'Infobip 短信发送失败', '重试3次，持续失败需人工介入'],
    ['AUX-6', 'EAI 接口超时', '系统告警，等待重试'],
    ['AUX-7', '文档自动挂载失败（补件场景）', '提示 document not exist，允许切换手动上传'],
    ['AUX-8', 'LHDN XML 生成失败', '记录失败日志，次日批次重试'],
], widths=[1.5, 5.0, 9.0])

doc.add_page_break()

# =====================================================================
# 19 待确认事项
# =====================================================================
H(1, '20. 待确认事项汇总')

H(2, '20.1 已结案项')
T(['编号', '事项', '结论'], [
    ['P-04', '指纹重试次数默认值', '第1次3次、第2次3次，由 BIOMETRIC_ATTEMPT_COUNTER 配置，系统上限999，可配置'],
    ['P-07', 'Amendment 变更追踪的完整字段清单', '已从原始 FRS 补齐，共46个字段，见 8.5 节'],
    ['P-11', 'CILT Appeal 流程中的 E-Acceptance 处理逻辑', '已明确，见 8.1 节；首次为 Manual 时免重签见 8.7 节'],
    ['P-17', 'OTP 最大错误输入次数', '5次'],
    ['P-22', 'Letter Template 完整清单', '已从原始 FRS 补齐，共8个模板及字段替代规则，见第15章'],
    ['R-A', '人工签与电子签的切换次数与条件', '2026-09-01 评审已明确，见第4章切换矩阵'],
    ['R-B', 'CILT 界面的操作步骤设计', '2026-09-01 评审确定合并为两步，见 8.3 节'],
], widths=[1.5, 5.5, 8.5])

H(2, '20.2 待客户/业务确认项')
T(['编号', '事项', '影响范围', '优先级'], [
    ['P-33', '🔴 RIB → CrediOS 的 OTP 验证结果回传接口（当前缺失）', '不解决则第2次验证失败后流程停滞', '高（阻塞）'],
    ['P-35', '🔴 带电子签名的文书样本；是否需符合规范的 PDF 数字签名/电子印章', '文书生成方案与行内签章能力', '高（阻塞）'],
    ['P-29', 'E_ACCEPTANCE_COMBO 生产取值（E[E] 还是 E[M]）', '决定所有「转 Manual」降级路径是否可行', '高'],
    ['P-32', '新状态模型与 e-Tracker 四个指示器的映射关系', 'e-Tracker 进度展示正确性', '高'],
    ['P-30', '上线时存量/在途案件的处理方式', '切换方案与数据迁移', '高'],
    ['P-36', '生物识别设备型号清单、各自返回码及含义', '设备联调', '中'],
    ['P-37', '设备返回码映射由前端还是后端实现？前端能否识别设备型号', '技术方案', '中'],
    ['P-38', 'CRA Report 四字段取值规则（尤其 Agreement Acceptance Type 与前序字段是否冲突）', '报表口径', '中'],
    ['P-39', '第2次生物识别的超时由 CrediOS 还是 RIB 控制', '超时处理归属', '中'],
    ['P-34', '文件中心「软覆盖」的技术实现方式', '归档方案', '中'],
    ['P-23', 'Owner 角色是否保留？企业客群车主是否需生物识别', '准入矩阵设计', '中'],
    ['P-24', '卖家（Seller Direct）生物识别功能是否保留', '功能范围', '中'],
    ['P-25', '外籍客户替代材料的具体要求', '豁免路径完整性', '中'],
    ['P-26', '列表页失败态展示文本方案', '运营分拣效率', '中'],
    ['P-27', '生物识别独立菜单的具体位置', '信息架构', '中'],
    ['P-28', 'Print Log 明细字段最终方案', '报表设计', '中'],
    ['P-31', 'LHDN XML 生成方式与报文结构是否变化', '合规产物', '中'],
    ['P-01', '纸质签字时效限制', '流程时效管理', '低'],
    ['P-02', '签约流程总时长限制', '流程时效管理', '低'],
    ['P-03', '销售发起后客户完成时限', '流程时效管理', '低'],
    ['P-05', '设备异常/网络中断的详细处理流程', '异常处理完整性', '低'],
    ['P-06', '冲正/撤销流程（签约完成后是否可撤销）', '流程完整性', '低'],
    ['P-08', '修订字段高亮的具体样式规范', 'UI 规范', '低'],
    ['P-09', 'RIB ↔ CrediOS 接口协议细节', '需索取 TRS 文档', '低'],
    ['P-10', '电子签约涉及的所有短信模板', '通知完整性', '低'],
    ['P-12', '附件大小限制具体值', '上传约束', '低'],
    ['P-14', 'e-Tracker General Constant 实际配置值', '参数配置', '低'],
    ['P-15', 'AEM e-Tracker Field Mapping 完整映射', '接口对接', '低'],
    ['P-16', '客户未注册 RIB 时的引导流程', '异常处理', '低'],
    ['P-18', 'OTP 重新发送冷却时间与每日上限', '防滥用', '低'],
], widths=[1.5, 6.5, 5.0, 2.5])

doc.save(r'F:\国际业务部\HLB\需求\个人签约\E-Acceptance需求20260901_v3.docx')
print('SAVED')
